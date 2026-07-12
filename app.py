"""
VineGuard AI - Sistema de Diagnóstico de Enfermedades en Uvas
Versión optimizada con Pruebas Estadísticas (Matthews y McNemar) + Multiidioma
"""

import sys
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')
elif hasattr(sys.stdout, 'buffer'):
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import streamlit as st
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.backends.backend_pdf import PdfPages
from PIL import Image
import tensorflow as tf
try:
    from tensorflow.keras.preprocessing.image import img_to_array
except ImportError:
    def img_to_array(img, **kwargs):
        return np.array(img, dtype=np.float32)
import os
import sys
import time
from pathlib import Path
from datetime import datetime
import pandas as pd
import base64
from scipy import stats
from sklearn.metrics import matthews_corrcoef, confusion_matrix
import tempfile
try:
    import joblib
    HAS_JOBLIB = True
except ImportError:
    HAS_JOBLIB = False
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    HAS_XLSX = True
except ImportError:
    HAS_XLSX = False

# Agregar src/ al path para importar extract_features
_SRC_PATH = Path(__file__).resolve().parent / "src"
if str(_SRC_PATH) not in sys.path:
    sys.path.insert(0, str(_SRC_PATH))

try:
    from extract_features import extract_single_image_features
    from mantenedor import SVM_SCALER_PATH, KNN_SCALER_PATH
    HAS_EXTRACT = True
except ImportError:
    HAS_EXTRACT = False
    SVM_SCALER_PATH = None
    KNN_SCALER_PATH = None

# ======= CONFIGURACIÓN MULTIIDIOMA =======
TRANSLATIONS = {
    'es': {
        'title': '🍇 VineGuard AI',
        'subtitle': 'Sistema Inteligente de Diagnóstico de Enfermedades en Viñedos',
        'subtitle_analysis': 'Con Análisis Estadístico Avanzado (Matthews & McNemar)',
        'language_selector': 'Idioma / Language',
        'config_title': '⚙️ Configuración',
        'load_models': '🚀 Cargar Modelos',
        'models_ready': '✅ Modelos listos',
        'available_models': '📊 Modelos Disponibles',
        'best_model': '🏆 Mejor Modelo',
        'model_ranking': '📊 Ranking de Modelos',
        'info_title': 'ℹ️ Información',
        'info_description': '''Esta aplicación utiliza modelos de deep learning para detectar enfermedades en hojas de vid:
        
        • **Podredumbre Negra**
        • **Esca** 
        • **Tizón de la Hoja**
        • **Hojas Sanas**
        
        **Análisis Estadístico:**
        • Coeficiente de Matthews (con múltiples imágenes)
        • Prueba de McNemar (con múltiples imágenes)
        
        **💡 Tip:** Use la pestaña 'Validación McNemar' para análisis estadístico completo con su propio dataset.''',
        'load_models_sidebar': '👈 Por favor, carga los modelos desde la barra lateral',
        'tab_diagnosis': '🔍 Diagnóstico',
        'tab_statistical': '📊 Análisis Estadístico', 
        'tab_validation': '🔬 Validación McNemar',
        'tab_info': '📚 Información',
        'diagnosis_title': '🔍 Diagnóstico de Enfermedades',
        'input_method': 'Selecciona método de entrada:',
        'upload_image': '📷 Subir imagen',
        'use_camera': '📸 Usar cámara',
        'select_image': 'Selecciona una imagen de hoja de vid',
        'supported_formats': 'Formatos soportados: JPG, JPEG, PNG',
        'image_loaded': 'Imagen cargada',
        'analyze_image': '🔬 Analizar Imagen',
        'analyzing': 'Analizando imagen...',
        'analysis_completed': '✅ Análisis completado!',
        'diagnosis_results': '📋 Resultados del Diagnóstico',
        'confidence': 'confianza',
        'consensus_diagnosis': '🤝 Diagnóstico Consensuado',
        'final_diagnosis': 'Diagnóstico Final:',
        'coincidence': 'Coincidencia',
        'probability_distribution': '📊 Distribución de Probabilidades',
        'treatment_recommendations': '💡 Recomendaciones de Tratamiento',
        'severity': 'Gravedad:',
        'recommended_treatment': '🏥 Tratamiento Recomendado',
        'preventive_measures': '🛡️ Medidas Preventivas',
        'generate_report': '📄 Generar Reporte',
        'download_pdf': '📥 Descargar Reporte PDF',
        'download_word': '📥 Descargar Reporte Word',
        'download_excel': '📥 Descargar Reporte Excel',
        'generating_report': 'Generando reporte...',
        'download_pdf_button': '💾 Descargar PDF',
        'download_word_button': '💾 Descargar Word',
        'download_excel_button': '💾 Descargar Excel',
        'camera_info': '📸 La función de cámara requiere acceso al hardware del dispositivo',
        'camera_warning': 'Por favor, usa la opción de subir imagen por ahora',
        'login_title': '🔐 Iniciar Sesión',
        'username': 'Usuario',
        'password': 'Contraseña',
        'login_button': 'Entrar',
        'login_error': 'Usuario o contraseña incorrectos',
        'logout': 'Cerrar Sesión',
        'disease_classes': {
            'Black_rot': 'Podredumbre Negra',
            'Esca': 'Esca (Sarampión Negro)', 
            'Healthy': 'Sana',
            'Leaf_blight': 'Tizón de la Hoja'
        },
        'detailed_recommendations': 'Recomendaciones Detalladas',
        'for_diagnosis': 'Para el diagnóstico',
        'additional_info': 'Información Adicional',
        'consult_specialist': 'Consulte con un especialista en viticultura',
        'follow_treatment_schedule': 'Siga un calendario regular de tratamientos',
        'monitor_evolution': 'Monitoree la evolución de la enfermedad',
        'document_treatments': 'Documente todos los tratamientos aplicados',
        'no_specific_recommendations': 'No hay recomendaciones específicas disponibles'
    },
    'en': {
        'title': '🍇 VineGuard AI',
        'subtitle': 'Intelligent Vineyard Disease Diagnosis System',
        'subtitle_analysis': 'With Advanced Statistical Analysis (Matthews & McNemar)',
        'language_selector': 'Language / Idioma',
        'config_title': '⚙️ Configuration',
        'load_models': '🚀 Load Models',
        'models_ready': '✅ Models ready',
        'available_models': '📊 Available Models',
        'best_model': '🏆 Best Model',
        'model_ranking': '📊 Model Ranking',
        'info_title': 'ℹ️ Information',
        'info_description': '''This application uses deep learning models to detect diseases in vine leaves:
        
        • **Black Rot**
        • **Esca** 
        • **Leaf Blight**
        • **Healthy Leaves**
        
        **Statistical Analysis:**
        • Matthews Coefficient (with multiple images)
        • McNemar Test (with multiple images)
        
        **💡 Tip:** Use the 'McNemar Validation' tab for complete statistical analysis with your own dataset.''',
        'load_models_sidebar': '👈 Please load the models from the sidebar',
        'tab_diagnosis': '🔍 Diagnosis',
        'tab_statistical': '📊 Statistical Analysis',
        'tab_validation': '🔬 McNemar Validation',
        'tab_info': '📚 Information',
        'diagnosis_title': '🔍 Disease Diagnosis',
        'input_method': 'Select input method:',
        'upload_image': '📷 Upload image',
        'use_camera': '📸 Use camera',
        'select_image': 'Select a vine leaf image',
        'supported_formats': 'Supported formats: JPG, JPEG, PNG',
        'image_loaded': 'Image loaded',
        'analyze_image': '🔬 Analyze Image',
        'analyzing': 'Analyzing image...',
        'analysis_completed': '✅ Analysis completed!',
        'diagnosis_results': '📋 Diagnosis Results',
        'confidence': 'confidence',
        'consensus_diagnosis': '🤝 Consensus Diagnosis',
        'final_diagnosis': 'Final Diagnosis:',
        'coincidence': 'Agreement',
        'probability_distribution': '📊 Probability Distribution',
        'treatment_recommendations': '💡 Treatment Recommendations',
        'severity': 'Severity:',
        'recommended_treatment': '🏥 Recommended Treatment',
        'preventive_measures': '🛡️ Preventive Measures',
        'generate_report': '📄 Generate Report',
        'download_pdf': '📥 Download PDF Report',
        'download_word': '📥 Download Word Report',
        'download_excel': '📥 Download Excel Report',
        'generating_report': 'Generating report...',
        'download_pdf_button': '💾 Download PDF',
        'download_word_button': '💾 Download Word',
        'download_excel_button': '💾 Download Excel',
        'camera_info': '📸 Camera function requires device hardware access',
        'camera_warning': 'Please use the upload image option for now',
        'login_title': '🔐 Login',
        'username': 'Username',
        'password': 'Password',
        'login_button': 'Login',
        'login_error': 'Incorrect username or password',
        'logout': 'Logout',
        'disease_classes': {
            'Black_rot': 'Black Rot',
            'Esca': 'Esca (Black Measles)', 
            'Healthy': 'Healthy',
            'Leaf_blight': 'Leaf Blight'
        },
        'detailed_recommendations': 'Detailed Recommendations',
        'for_diagnosis': 'For diagnosis',
        'additional_info': 'Additional Information',
        'consult_specialist': 'Consult with a viticulture specialist',
        'follow_treatment_schedule': 'Follow a regular treatment schedule',
        'monitor_evolution': 'Monitor disease evolution',
        'document_treatments': 'Document all applied treatments',
        'no_specific_recommendations': 'No specific recommendations available'
    },
    'pt': {
        'title': '🍇 VineGuard AI',
        'subtitle': 'Sistema Inteligente de Diagnóstico de Doenças em Vinhedos',
        'subtitle_analysis': 'Com Análise Estatística Avançada (Matthews & McNemar)',
        'language_selector': 'Idioma / Language',
        'config_title': '⚙️ Configuração',
        'load_models': '🚀 Carregar Modelos',
        'models_ready': '✅ Modelos prontos',
        'available_models': '📊 Modelos Disponíveis',
        'best_model': '🏆 Melhor Modelo',
        'model_ranking': '📊 Ranking de Modelos',
        'info_title': 'ℹ️ Informação',
        'info_description': '''Esta aplicação usa modelos de deep learning para detectar doenças em folhas de videira:
        
        • **Podridão Negra**
        • **Esca** 
        • **Queima das Folhas**
        • **Folhas Saudáveis**
        
        **Análise Estatística:**
        • Coeficiente de Matthews (com múltiplas imagens)
        • Teste de McNemar (com múltiplas imagens)
        
        **💡 Dica:** Use a aba 'Validação McNemar' para análise estatística completa com seu próprio dataset.''',
        'load_models_sidebar': '👈 Por favor, carregue os modelos da barra lateral',
        'tab_diagnosis': '🔍 Diagnóstico',
        'tab_statistical': '📊 Análise Estatística',
        'tab_validation': '🔬 Validação McNemar',
        'tab_info': '📚 Informação',
        'diagnosis_title': '🔍 Diagnóstico de Doenças',
        'input_method': 'Selecione o método de entrada:',
        'upload_image': '📷 Carregar imagem',
        'use_camera': '📸 Usar câmera',
        'select_image': 'Selecione uma imagem de folha de videira',
        'supported_formats': 'Formatos suportados: JPG, JPEG, PNG',
        'image_loaded': 'Imagem carregada',
        'analyze_image': '🔬 Analisar Imagem',
        'analyzing': 'Analisando imagem...',
        'analysis_completed': '✅ Análise concluída!',
        'diagnosis_results': '📋 Resultados do Diagnóstico',
        'confidence': 'confiança',
        'consensus_diagnosis': '🤝 Diagnóstico Consensual',
        'final_diagnosis': 'Diagnóstico Final:',
        'coincidence': 'Concordância',
        'probability_distribution': '📊 Distribuição de Probabilidade',
        'treatment_recommendations': '💡 Recomendações de Tratamento',
        'severity': 'Gravidade:',
        'recommended_treatment': '🏥 Tratamento Recomendado',
        'preventive_measures': '🛡️ Medidas Preventivas',
        'generate_report': '📄 Gerar Relatório',
        'download_pdf': '📥 Baixar Relatório PDF',
        'download_word': '📥 Baixar Relatório Word',
        'download_excel': '📥 Baixar Relatório Excel',
        'generating_report': 'Gerando relatório...',
        'download_pdf_button': '💾 Baixar PDF',
        'download_word_button': '💾 Baixar Word',
        'download_excel_button': '💾 Baixar Excel',
        'camera_info': '📸 A função da câmera requer acesso ao hardware do dispositivo',
        'camera_warning': 'Por favor, use a opção de carregar imagem por enquanto',
        'login_title': '🔐 Iniciar Sessão',
        'username': 'Usuário',
        'password': 'Senha',
        'login_button': 'Entrar',
        'login_error': 'Usuário ou senha incorretos',
        'logout': 'Sair',
        'disease_classes': {
            'Black_rot': 'Podridão Negra',
            'Esca': 'Esca (Sarampo Negro)', 
            'Healthy': 'Saudável',
            'Leaf_blight': 'Queima das Folhas'
        },
        'detailed_recommendations': 'Recomendações Detalhadas',
        'for_diagnosis': 'Para o diagnóstico',
        'additional_info': 'Informações Adicionais',
        'consult_specialist': 'Consulte um especialista em viticultura',
        'follow_treatment_schedule': 'Siga um cronograma regular de tratamentos',
        'monitor_evolution': 'Monitore a evolução da doença',
        'document_treatments': 'Documente todos os tratamentos aplicados',
        'no_specific_recommendations': 'Não há recomendações específicas disponíveis'
    }
}

# Función para obtener texto traducido
def get_text(key, language='es'):
    """Obtiene texto traducido según el idioma seleccionado"""
    try:
        return TRANSLATIONS[language][key]
    except KeyError:
        # Fallback a español si la clave no existe
        return TRANSLATIONS['es'].get(key, key)

# Inicializar idioma en session_state
if 'language' not in st.session_state:
    st.session_state.language = 'es'

# Configuración de la página
st.set_page_config(
    page_title="VineGuard AI",
    page_icon="🍇",
    layout="centered",
    initial_sidebar_state="expanded"
)

# ======= SELECTOR DE IDIOMA EN LA PARTE SUPERIOR =======
col_lang1, col_lang2 = st.columns([3.5, 1.5])
with col_lang2:
    language_options = {
        '🇪🇸 Español': 'es',
        '🇺🇸 English': 'en', 
        '🇧🇷 Português': 'pt'
    }
    
    selected_language = st.selectbox(
        "🌐 Language / Idioma",
        options=list(language_options.keys()),
        index=list(language_options.values()).index(st.session_state.language),
        key="main_language_selector"
    )
    
    # Actualizar idioma si cambió
    new_language = language_options[selected_language]
    if new_language != st.session_state.language:
        st.session_state.language = new_language
        st.rerun()

st.markdown("---")

# CSS personalizado
st.markdown("""
<style>
    /* Diseño responsive */
    .main .block-container {
        padding: 1.5rem;
        max-width: 900px;
    }
    
    /* Botones grandes para móviles y consistencia visual */
    .stButton button {
        width: 100%;
        padding: 0.75rem;
        font-size: 1rem;
        background: linear-gradient(135deg, #2e7d32 0%, #4a148c 100%);
        color: white !important;
        border-radius: 10px;
        border: none;
        transition: transform 0.2s ease, box-shadow 0.2s ease;
        font-weight: 600;
    }
    
    .stButton button:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(74, 20, 140, 0.3);
        color: white !important;
    }
    
    /* Mejoras visuales */
    .stAlert {
        padding: 1.2rem;
        border-radius: 12px;
    }
    
    /* Ocultar elementos innecesarios */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    
    /* Estilo para métricas */
    [data-testid="metric-container"] {
        background-color: #f5f6f9;
        border: 1px solid #e2e8f0;
        padding: 12px;
        border-radius: 12px;
        margin: 6px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.02);
    }
    
    /* Estilo para estadísticas */
    .statistical-box {
        background-color: #f0fdf4;
        border: 2px solid #2e7d32;
        padding: 15px;
        border-radius: 12px;
        margin: 10px 0;
    }
    
    /* Bloques de diseño */
    .hero-container {
        background: linear-gradient(135deg, #e8f5e9 0%, #f3e5f5 100%);
        border-radius: 20px;
        padding: 30px;
        margin-bottom: 25px;
        border: 1px solid #d1c4e9;
        box-shadow: 0 4px 20px rgba(0, 0, 0, 0.05);
        text-align: center;
    }
    
    .hero-title {
        color: #4a148c !important;
        font-size: 2.5rem !important;
        font-weight: 800 !important;
        margin: 0 0 5px 0 !important;
    }
    
    .hero-subtitle {
        color: #2e7d32 !important;
        font-size: 1.25rem !important;
        font-weight: 500 !important;
        margin: 0 0 15px 0 !important;
        line-height: 1.4 !important;
    }
    
    .hero-description {
        color: #424242 !important;
        font-size: 1rem !important;
        line-height: 1.6 !important;
        max-width: 700px;
        margin: 0 auto !important;
    }
    
    /* Tarjetas de características */
    .features-grid {
        display: grid;
        grid-template-columns: 1fr 1fr;
        gap: 15px;
        margin-top: 25px;
    }
    
    @media (max-width: 768px) {
        .features-grid {
            grid-template-columns: 1fr;
        }
    }
    
    .feature-card {
        background-color: white;
        border-radius: 15px;
        padding: 20px;
        border: 1px solid #e0e0e0;
        box-shadow: 0 4px 12px rgba(0,0,0,0.03);
        transition: transform 0.3s ease;
        display: flex;
        align-items: flex-start;
        text-align: left;
    }
    
    .feature-card:hover {
        transform: translateY(-3px);
        box-shadow: 0 8px 16px rgba(0,0,0,0.06);
        border-color: #b39ddb;
    }
    
    .feature-icon {
        font-size: 1.8rem;
        margin-right: 15px;
        line-height: 1;
    }
    
    .feature-content h4 {
        margin: 0 0 5px 0 !important;
        color: #4a148c !important;
        font-size: 1.1rem !important;
        font-weight: 600 !important;
    }
    
    .feature-content p {
        margin: 0 !important;
        color: #616161 !important;
        font-size: 0.9rem !important;
        line-height: 1.4 !important;
    }
    
    /* Tarjeta de instrucciones antes de cargar modelos */
    .load-warning-card {
        background: #fffdf5;
        border-left: 5px solid #ffb300;
        padding: 20px;
        border-radius: 12px;
        box-shadow: 0 4px 12px rgba(0,0,0,0.02);
        margin: 20px 0;
        text-align: left;
    }
    
    .load-warning-title {
        color: #ff8f00 !important;
        margin: 0 0 8px 0 !important;
        font-weight: 700 !important;
        font-size: 1.15rem !important;
    }
    
    .load-warning-text {
        color: #5d4037 !important;
        margin: 0 !important;
        font-size: 0.95rem !important;
        line-height: 1.5 !important;
    }
    
    .load-warning-note {
        color: #795548 !important;
        margin: 8px 0 0 0 !important;
        font-size: 0.85rem !important;
        font-style: italic !important;
    }

    /* Información de barra lateral */
    .sidebar-info-block {
        background-color: #f9f9fb;
        border-radius: 12px;
        padding: 15px;
        border: 1px solid #edeef2;
        margin-top: 15px;
        text-align: left;
    }
    
    .sidebar-info-title {
        font-size: 0.95rem !important;
        font-weight: 700 !important;
        color: #4a148c !important;
        margin: 0 0 10px 0 !important;
    }

    .sidebar-info-item {
        font-size: 0.85rem !important;
        color: #424242 !important;
        margin-bottom: 6px !important;
        line-height: 1.4 !important;
    }
    
    .sidebar-info-label {
        font-weight: 600;
        color: #2e7d32;
    }
    
    .status-badge {
        display: inline-block;
        padding: 2px 8px;
        border-radius: 20px;
        font-size: 0.75rem;
        font-weight: 700;
        text-transform: uppercase;
    }
    
    .status-badge.loaded {
        background-color: #e8f5e9;
        color: #2e7d32;
    }
    
    .status-badge.not-loaded {
        background-color: #ffebee;
        color: #c62828;
    }
    
    /* Estilo para cajas de teoría */
    .theory-box {
        background: linear-gradient(135deg, #4a148c 0%, #2e7d32 100%);
        padding: 20px;
        border-radius: 15px;
        color: white;
        margin: 10px 0;
        box-shadow: 0 4px 15px rgba(0,0,0,0.1);
    }
    
    .theory-box h4 {
        color: white !important;
        margin-bottom: 10px;
    }
    
    .theory-box p {
        color: #f0f0f0 !important;
        line-height: 1.6;
    }
    
    /* Estilo para carpetas de enfermedades */
    .disease-folder {
        background: linear-gradient(135deg, #ff9a9e 0%, #fecfef 50%, #fecfef 100%);
        padding: 20px;
        border-radius: 15px;
        margin: 10px 0;
        box-shadow: 0 4px 15px rgba(0,0,0,0.1);
        border: 2px solid #ff6b6b;
    }
    
    .disease-folder.black-rot {
        background: linear-gradient(135deg, #ff6b6b 0%, #ff8e8e 100%);
        border-color: #dc3545;
    }
    
    .disease-folder.esca {
        background: linear-gradient(135deg, #8B4513 0%, #CD853F 100%);
        border-color: #8B4513;
    }
    
    .disease-folder.healthy {
        background: linear-gradient(135deg, #28a745 0%, #20c997 100%);
        border-color: #28a745;
    }
    
    .disease-folder.leaf-blight {
        background: linear-gradient(135deg, #ffc107 0%, #ffeb3b 100%);
        border-color: #ffc107;
    }
    
    /* Resultados destacados */
    .result-highlight {
        background: linear-gradient(135deg, #4a148c 0%, #2e7d32 100%);
        padding: 20px;
        border-radius: 15px;
        color: white;
        margin: 15px 0;
        text-align: center;
        box-shadow: 0 4px 15px rgba(0,0,0,0.2);
    }
    
    .interpretation-box {
        background: linear-gradient(135deg, #11998e 0%, #38ef7d 100%);
        padding: 25px;
        border-radius: 15px;
        color: white;
        margin: 20px 0;
        box-shadow: 0 6px 20px rgba(0,0,0,0.15);
    }
    
    .interpretation-box h3 {
        color: white !important;
        margin-bottom: 15px;
    }
    
    .interpretation-box p {
        color: #f0f0f0 !important;
        font-size: 1.1em;
        line-height: 1.7;
    }
</style>
""", unsafe_allow_html=True)

# Configuración de los 5 modelos: 3 clásicos + 2 híbridos
MODELS_CONFIG = {
    "M1 - SVM": {
        "type": "classic",
        "path": "models/svm_model.pkl",
        "description": "Support Vector Machine (RBF) + características manuales",
    },
    "M2 - Random Forest": {
        "type": "classic",
        "path": "models/random_forest_model.pkl",
        "description": "Random Forest (200 árboles) + características manuales",
    },
    "M3 - KNN": {
        "type": "classic",
        "path": "models/knn_model.pkl",
        "description": "K-Nearest Neighbors (k=5) + características manuales",
    },
    "H1 - CNN + SVM": {
        "type": "hybrid_cnn_svm",
        "extractor_path": "models/cnn_feature_extractor.h5",
        "classifier_path": "models/cnn_svm_model.pkl",
        "description": "CNN extractor de features + clasificador SVM",
    },
    "H2 - Transfer + RF": {
        "type": "hybrid_transfer_rf",
        "extractor_path": "models/transfer_feature_extractor.h5",
        "classifier_path": "models/transfer_random_forest_model.pkl",
        "description": "MobileNetV2 (ImageNet) + Random Forest",
    },
}

# Mantener MODEL_PATHS para compatibilidad con partes del código que lo referencian
MODEL_PATHS = {k: v.get("path", v.get("extractor_path", "")) for k, v in MODELS_CONFIG.items()}

# Clases de enfermedades (keys en inglés para consistencia)
DISEASE_CLASSES = ["Black_rot", "Esca", "Healthy", "Leaf_blight"]

# Función para obtener nombres de enfermedades según idioma
def get_disease_names(language='es'):
    """Retorna diccionario de nombres de enfermedades según idioma"""
    return get_text('disease_classes', language)

# Función para obtener configuración de carpetas según idioma
def get_disease_folders(language='es'):
    """Retorna configuración de carpetas según idioma"""
    disease_names = get_disease_names(language)
    
    if language == 'en':
        return {
            disease_names["Black_rot"]: {
                "key": "Black_rot",
                "icon": "🔴",
                "description": "Guignardia bidwellii fungi",
                "css_class": "black-rot"
            },
            disease_names["Esca"]: {
                "key": "Esca",
                "icon": "🟤", 
                "description": "Vascular fungi complex",
                "css_class": "esca"
            },
            f"{disease_names['Healthy']} Leaves": {
                "key": "Healthy",
                "icon": "✅",
                "description": "No detectable diseases",
                "css_class": "healthy"
            },
            disease_names["Leaf_blight"]: {
                "key": "Leaf_blight",
                "icon": "🟡",
                "description": "Isariopsis fungi",
                "css_class": "leaf-blight"
            }
        }
    elif language == 'pt':
        return {
            disease_names["Black_rot"]: {
                "key": "Black_rot",
                "icon": "🔴",
                "description": "Fungos Guignardia bidwellii",
                "css_class": "black-rot"
            },
            disease_names["Esca"]: {
                "key": "Esca",
                "icon": "🟤",
                "description": "Complexo de fungos vasculares",
                "css_class": "esca"
            },
            f"Folhas {disease_names['Healthy']}": {
                "key": "Healthy",
                "icon": "✅",
                "description": "Sem doenças detectáveis",
                "css_class": "healthy"
            },
            disease_names["Leaf_blight"]: {
                "key": "Leaf_blight",
                "icon": "🟡",
                "description": "Fungo Isariopsis",
                "css_class": "leaf-blight"
            }
        }
    else:  # Español por defecto
        return {
            disease_names["Black_rot"]: {
                "key": "Black_rot",
                "icon": "🔴",
                "description": "Hongos Guignardia bidwellii",
                "css_class": "black-rot"
            },
            disease_names["Esca"]: {
                "key": "Esca",
                "icon": "🟤",
                "description": "Complejo de hongos vasculares",
                "css_class": "esca"
            },
            f"Hojas {disease_names['Healthy']}": {
                "key": "Healthy",
                "icon": "✅",
                "description": "Sin enfermedades detectables",
                "css_class": "healthy"
            },
            disease_names["Leaf_blight"]: {
                "key": "Leaf_blight",
                "icon": "🟡",
                "description": "Hongo Isariopsis",
                "css_class": "leaf-blight"
            }
        }

# Función para cargar ranking de modelos
def load_model_ranking():
    ranking_path = Path("reports/modelos/ranking_modelos.csv")
    best_path = Path("reports/modelos/mejor_modelo.txt")
    if ranking_path.exists():
        df = pd.read_csv(ranking_path)
        ranking_data = df.to_dict('records')
    else:
        ranking_data = None
    best_model_name = None
    if best_path.exists():
        with open(best_path, 'r', encoding='utf-8') as f:
            best_model_name = f.read().strip()
    return ranking_data, best_model_name

def check_pipeline_status():
    checks = {
        "eda": {
            "label_es": "EDA - Análisis exploratorio",
            "label_en": "EDA - Exploratory analysis",
            "label_pt": "EDA - Análise exploratória",
            "files": [Path("reports/eda/resumen_dataset.csv"), Path("reports/eda/distribucion_clases.png")],
        },
        "preprocessing": {
            "label_es": "Preprocesamiento y aumento",
            "label_en": "Preprocessing & augmentation",
            "label_pt": "Pré-processamento e aumento",
            "files": [Path("reports/preprocessing/ejemplos_aumento_datos.png")],
        },
        "crossval": {
            "label_es": "Validación cruzada (5-folds)",
            "label_en": "Cross-validation (5-fold)",
            "label_pt": "Validação cruzada (5-fold)",
            "files": [Path("reports/modelos/cross_validation_resultados.csv")],
        },
        "hyperparam": {
            "label_es": "Optimización de hiperparámetros",
            "label_en": "Hyperparameter optimization",
            "label_pt": "Otimização de hiperparâmetros",
            "files": [Path("reports/modelos/mejores_hiperparametros.csv")],
        },
        "statistical": {
            "label_es": "Validación estadística",
            "label_en": "Statistical validation",
            "label_pt": "Validação estatística",
            "files": [Path("reports/estadistica/mcnemar_resultados.csv")],
        },
        "model_selection": {
            "label_es": "Selección del mejor modelo",
            "label_en": "Best model selection",
            "label_pt": "Seleção do melhor modelo",
            "files": [Path("reports/modelos/ranking_modelos.csv")],
        },
    }
    lang = st.session_state.get("language", "es")
    results = {}
    for key, check in checks.items():
        done = all(f.exists() for f in check["files"])
        label = check.get(f"label_{lang}", check["label_es"])
        results[key] = {"done": done, "label": label}
    return results

# Inicializar estado de sesión
if 'logged_in' not in st.session_state:
    st.session_state.logged_in = False
if 'models_loaded' not in st.session_state:
    st.session_state.models_loaded = False
    st.session_state.models = {}
    st.session_state.current_image = None
    st.session_state.predictions = None
    st.session_state.statistical_analysis = None
    st.session_state.mcnemar_validation = None
    st.session_state.mcnemar_analysis = None
    st.session_state.ranking_data, st.session_state.best_model_name = load_model_ranking()
    st.session_state.pdf_bytes = None
    st.session_state.pdf_ready = False
    st.session_state.docx_bytes = None
    st.session_state.docx_ready = False
    st.session_state.xlsx_bytes = None
    st.session_state.xlsx_ready = False

# Credenciales de usuario (puedes modificar estas)
VALID_CREDENTIALS = {
    "admin": "admin123",
    "usuario": "12345"
}

# Mostrar pantalla de login si no está autenticado
if not st.session_state.logged_in:
    st.markdown("<h1 style='text-align: center; color: #4a148c;'>" + get_text('title', st.session_state.language) + "</h1>", unsafe_allow_html=True)
    st.markdown("<h2 style='text-align: center; color: #2e7d32;'>" + get_text('login_title', st.session_state.language) + "</h2>", unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        username = st.text_input(get_text('username', st.session_state.language))
        password = st.text_input(get_text('password', st.session_state.language), type='password')
        if st.button(get_text('login_button', st.session_state.language), use_container_width=True):
            if username in VALID_CREDENTIALS and VALID_CREDENTIALS[username] == password:
                st.session_state.logged_in = True
                st.rerun()
            else:
                st.error(get_text('login_error', st.session_state.language))
    
    st.stop()

# Botón de cerrar sesión
if st.sidebar.button(get_text('logout', st.session_state.language)):
    st.session_state.logged_in = False
    st.rerun()

# ─── Carga de modelos ────────────────────────────────────────────────────────
@st.cache_resource
def load_models():
    """
    Carga los 5 modelos:
      - M1 SVM, M2 RF, M3 KNN  → joblib (.pkl)
      - H1 CNN+SVM              → TF extractor (.h5) + joblib SVM (.pkl)
      - H2 Transfer+RF          → TF extractor (.h5) + joblib RF (.pkl)
    Retorna dict con objetos listos para inferencia.
    """
    loaded = {}
    for name, cfg in MODELS_CONFIG.items():
        model_type = cfg["type"]
        try:
            if model_type == "classic":
                path = cfg["path"]
                if os.path.exists(path):
                    loaded[name] = {"type": model_type, "clf": joblib.load(path)}
                    print(f"✓ {name} cargado desde {path}")
                else:
                    st.warning(f"⚠️ Modelo no encontrado: {path}")

            elif model_type in ("hybrid_cnn_svm", "hybrid_transfer_rf"):
                ext_path = cfg["extractor_path"]
                clf_path = cfg["classifier_path"]
                if os.path.exists(ext_path) and os.path.exists(clf_path):
                    extractor = tf.keras.models.load_model(ext_path)
                    clf = joblib.load(clf_path)
                    loaded[name] = {
                        "type": model_type,
                        "extractor": extractor,
                        "clf": clf,
                    }
                    print(f"✓ {name} cargado")
                else:
                    missing = [p for p in [ext_path, clf_path] if not os.path.exists(p)]
                    st.warning(f"⚠️ Archivos no encontrados para {name}: {missing}")

        except Exception as e:
            st.error(f"Error al cargar {name}: {str(e)}")

    return loaded


# ─── Preprocesamiento por tipo de modelo ────────────────────────────────────

def preprocess_for_classic(image, model_name=None):
    """
    Para modelos clásicos (SVM, RF, KNN):
    Extrae características manuales con extract_features.py y aplica el scaler
    CORRESPONDIENTE al modelo (cada uno fue entrenado con su propio scaler,
    o sin scaler en el caso de Random Forest).
    Retorna array numpy (1, n_features).
    """
    if HAS_EXTRACT:
        if model_name == "M1 - SVM" and SVM_SCALER_PATH is not None:
            return extract_single_image_features(image, apply_scaler=True, scaler_path=SVM_SCALER_PATH)
        elif model_name == "M3 - KNN" and KNN_SCALER_PATH is not None:
            return extract_single_image_features(image, apply_scaler=True, scaler_path=KNN_SCALER_PATH)
        elif model_name == "M2 - Random Forest":
            return extract_single_image_features(image, apply_scaler=False)
        return extract_single_image_features(image, apply_scaler=False)
    # Fallback sin scikit-image: histograma + estadísticas RGB planas
    img = image.resize((224, 224))
    arr = np.array(img, dtype=np.float32)
    feats = []
    for c in range(3):
        ch = arr[:, :, c]
        feats.extend([ch.mean(), ch.std(), ch.var()])
        hist, _ = np.histogram(ch, bins=64, range=(0, 256))
        hist = hist.astype(np.float32)
        if hist.sum() > 0:
            hist /= hist.sum()
        feats.extend(hist.tolist())
    return np.array(feats, dtype=np.float32).reshape(1, -1)


def preprocess_for_cnn_extractor(image, target_size=(224, 224)):
    """
    Para H1 (CNN extractor): normaliza la imagen a [0, 1].
    Retorna tensor (1, H, W, 3).
    """
    img = image.resize(target_size)
    arr = np.array(img, dtype=np.float32) / 255.0
    return np.expand_dims(arr, axis=0)


def preprocess_for_transfer_extractor(image, target_size=(224, 224)):
    """
    Para H2 (MobileNetV2): aplica preprocess_input de MobileNetV2.
    Retorna tensor (1, H, W, 3) en rango [-1, 1].
    """
    from tensorflow.keras.applications.mobilenet_v2 import preprocess_input
    img = image.resize(target_size)
    arr = np.array(img, dtype=np.float32)
    arr = preprocess_input(arr)
    return np.expand_dims(arr, axis=0)


# Función legacy para mantener compatibilidad con código existente
def preprocess_image(image, target_size=(224, 224), model_name=None):
    """Preprocesa imagen (legacy — redirige al preprocesamiento apropiado)."""
    return preprocess_for_cnn_extractor(image, target_size)


# ─── Predicción universal ────────────────────────────────────────────────────
def predict_disease(image, model_bundle, model_name):
    """
    Realiza predicción con cualquiera de los 5 modelos.

    Parameters
    ----------
    image        : PIL.Image
    model_bundle : dict con claves 'type', 'clf' (y opcionalmente 'extractor')
    model_name   : str — nombre visible del modelo

    Returns
    -------
    dict con predicted_class, confidence, all_predictions, inference_time, etc.
    """
    model_type = model_bundle.get("type", "classic")
    start_time = time.time()

    if model_type == "classic":
        # ── M1, M2, M3: características manuales ───────────────────────────
        feats = preprocess_for_classic(image, model_name)   # (1, n_features)
        clf = model_bundle["clf"]
        proba = clf.predict_proba(feats)[0]             # (n_classes,)

    elif model_type == "hybrid_cnn_svm":
        # ── H1: CNN extractor → SVM ─────────────────────────────────────────
        img_tensor = preprocess_for_cnn_extractor(image)   # (1, 224, 224, 3)
        extractor = model_bundle["extractor"]
        feats = extractor(img_tensor, training=False).numpy()  # (1, 256)
        clf = model_bundle["clf"]
        proba = clf.predict_proba(feats)[0]             # (n_classes,)

    elif model_type == "hybrid_transfer_rf":
        # ── H2: MobileNetV2 extractor → RF ─────────────────────────────────
        img_tensor = preprocess_for_transfer_extractor(image)  # (1, 224, 224, 3)
        extractor = model_bundle["extractor"]
        feats = extractor(img_tensor, training=False).numpy()   # (1, 1280)
        clf = model_bundle["clf"]
        proba = clf.predict_proba(feats)[0]             # (n_classes,)

    else:
        raise ValueError(f"Tipo de modelo desconocido: {model_type}")

    inference_time = (time.time() - start_time) * 1000  # ms

    # Asegurar que proba tenga longitud = n_clases
    n_classes = len(DISEASE_CLASSES)
    if len(proba) < n_classes:
        padded = np.zeros(n_classes, dtype=np.float32)
        clf_classes = model_bundle["clf"].classes_
        for i, cls_idx in enumerate(clf_classes):
            if cls_idx < n_classes:
                padded[cls_idx] = proba[i]
        proba = padded

    predicted_class_idx = int(np.argmax(proba))
    predicted_class = DISEASE_CLASSES[predicted_class_idx]
    confidence = float(proba[predicted_class_idx])

    return {
        'model_name': model_name,
        'predicted_class': predicted_class,
        'predicted_class_es': get_disease_names(st.session_state.language)[predicted_class],
        'confidence': confidence,
        'all_predictions': proba,
        'inference_time': inference_time,
        'predicted_class_idx': predicted_class_idx,
    }

# ======= NUEVAS FUNCIONES ESTADÍSTICAS =======

def calculate_matthews_coefficient(y_true, y_pred, num_classes):
    """
    Calcula el Coeficiente de Matthews para clasificación multiclase
    """
    try:
        mcc = matthews_corrcoef(y_true, y_pred)
        return mcc
    except:
        # Cálculo manual si hay problemas
        cm = confusion_matrix(y_true, y_pred, labels=range(num_classes))

        # Para multiclase, usamos la fórmula generalizada
        # MCC = (∑c*s - ∑pk*tk) / sqrt((∑s^2 - ∑pk^2)(∑s^2 - ∑tk^2))

        n = cm.sum()
        sum_diag = np.trace(cm)

        sum_pk = np.sum(cm.sum(axis=0) ** 2)
        sum_tk = np.sum(cm.sum(axis=1) ** 2)
        sum_squares = np.sum(cm.sum(axis=0) * cm.sum(axis=1))

        numerator = n * sum_diag - sum_squares
        denominator = np.sqrt((n**2 - sum_pk) * (n**2 - sum_tk))

        if denominator == 0:
            return 0.0

        mcc = numerator / denominator
        return mcc

def mcnemar_test_multiclass(y_true, y_pred1, y_pred2):
    """
    Prueba de McNemar para clasificación multiclase
    Compara si dos modelos difieren significativamente en sus predicciones
    """
    # Crear tabla de contingencia 2x2
    # (correcto_modelo1, incorrecto_modelo1) vs (correcto_modelo2, incorrecto_modelo2)

    correct_1 = (y_true == y_pred1)
    correct_2 = (y_true == y_pred2)

    # Casos donde los modelos difieren
    model1_correct_model2_wrong = np.sum(correct_1 & ~correct_2)  # b
    model1_wrong_model2_correct = np.sum(~correct_1 & correct_2)  # c

    # Tabla de McNemar
    # |  Modelo2  |           |
    # |  C    W   | Modelo1   |
    # |  a    b   | Correcto  |
    # |  c    d   | Incorrecto|

    b = model1_correct_model2_wrong
    c = model1_wrong_model2_correct

    # Si no hay diferencias, no se puede hacer la prueba
    if b + c == 0:
        return {
            'statistic': 0.0,
            'p_value': 1.0,
            'b': b,
            'c': c,
            'interpretation': 'No hay diferencias entre modelos'
        }

    # Aplicar corrección de continuidad de Yates
    if b + c > 25:
        # Para muestras grandes, usar corrección de continuidad
        statistic = (abs(b - c) - 0.5) ** 2 / (b + c)
    else:
        # Para muestras pequeñas, usar prueba exacta
        statistic = (b - c) ** 2 / (b + c)

    # Calcular p-valor usando distribución chi-cuadrado con 1 grado de libertad
    p_value = 1 - stats.chi2.cdf(statistic, df=1)

    # Interpretación
    if p_value < 0.001:
        interpretation = "Diferencia altamente significativa (p < 0.001)"
    elif p_value < 0.01:
        interpretation = "Diferencia muy significativa (p < 0.01)"
    elif p_value < 0.05:
        interpretation = "Diferencia significativa (p < 0.05)"
    elif p_value < 0.1:
        interpretation = "Diferencia marginalmente significativa (p < 0.1)"
    else:
        interpretation = "No hay diferencia significativa (p ≥ 0.1)"

    return {
        'statistic': statistic,
        'p_value': p_value,
        'b': b,
        'c': c,
        'interpretation': interpretation
    }

def interpret_mcc(mcc):
    """Interpreta el valor del Coeficiente de Matthews"""
    if mcc >= 0.9:
        return "Excelente (≥ 0.9)"
    elif mcc >= 0.8:
        return "Muy bueno (0.8-0.89)"
    elif mcc >= 0.6:
        return "Bueno (0.6-0.79)"
    elif mcc >= 0.4:
        return "Moderado (0.4-0.59)"
    elif mcc >= 0.2:
        return "Débil (0.2-0.39)"
    elif mcc > 0:
        return "Muy débil (0-0.19)"
    elif mcc == 0:
        return "Sin correlación (0)"
    else:
        return "Correlación negativa (< 0)"

# ======= FUNCIONES PARA VALIDACIÓN CON MÚLTIPLES IMÁGENES =======

def process_multiple_images_by_folders(disease_files, models):
    """
    Procesa múltiples imágenes organizadas por carpetas de enfermedades
    """
    all_predictions = {model_name: [] for model_name in models.keys()}
    y_true = []
    total_images = 0

    # Contar total de imágenes
    for disease_name, files in disease_files.items():
        total_images += len(files)

    if total_images == 0:
        return None, "No se cargaron imágenes"

    try:
        progress_bar = st.progress(0)
        processed = 0

        for disease_name, files in disease_files.items():
            if len(files) > 0:
                # Obtener la clave en inglés de la enfermedad
                disease_folders = get_disease_folders(st.session_state.language)
                disease_key = disease_folders[disease_name]["key"]
                disease_idx = DISEASE_CLASSES.index(disease_key)

                for uploaded_file in files:
                    # Cargar imagen
                    image = Image.open(uploaded_file).convert('RGB')

                    # Añadir etiqueta verdadera
                    y_true.append(disease_idx)

                    # Obtener predicciones de todos los modelos
                    for model_name, model in models.items():
                        result = predict_disease(image, model, model_name)
                        predicted_idx = result['predicted_class_idx']
                        all_predictions[model_name].append(predicted_idx)

                    processed += 1
                    progress_bar.progress(processed / total_images)

        progress_bar.empty()

        # Convertir a arrays numpy
        model_predictions = [np.array(all_predictions[model_name]) for model_name in models.keys()]
        y_true = np.array(y_true)

        return {
            'y_true': y_true,
            'predictions': model_predictions,
            'model_names': list(models.keys())
        }, None

    except Exception as e:
        return None, f"Error procesando imágenes: {str(e)}"

def create_validation_results_display(validation_data, mcnemar_analysis):
    """
    Crea visualización de resultados de validación
    """
    y_true = validation_data['y_true']
    model_predictions = validation_data['predictions']
    model_names = validation_data['model_names']

    # Calcular métricas por modelo
    results_summary = []
    for i, (model_name, predictions) in enumerate(zip(model_names, model_predictions)):
        accuracy = np.mean(y_true == predictions)
        results_summary.append({
            'Modelo': model_name,
            'Precisión': f"{accuracy:.1%}",
            'Muestras Correctas': f"{np.sum(y_true == predictions)}/{len(y_true)}"
        })

    return pd.DataFrame(results_summary)

def perform_mcnemar_analysis(validation_data):
    """
    Realiza análisis McNemar con datos reales de validación
    """
    if validation_data is None:
        return None

    y_true_real = validation_data['y_true']
    model_predictions = validation_data['predictions']
    model_names = validation_data['model_names']

    # Calcular MCC real para cada modelo
    matthews_coefficients = []
    for i, (model_name, predictions) in enumerate(zip(model_names, model_predictions)):
        mcc = calculate_matthews_coefficient(y_true_real, predictions, len(DISEASE_CLASSES))
        matthews_coefficients.append({
            'model': model_name,
            'mcc': mcc,
            'interpretation': interpret_mcc(mcc)
        })

    # Realizar pruebas de McNemar entre todos los pares de modelos
    mcnemar_results = []
    for i in range(len(model_names)):
        for j in range(i + 1, len(model_names)):
            if i < len(model_predictions) and j < len(model_predictions):
                mcnemar_result = mcnemar_test_multiclass(
                    y_true_real,
                    model_predictions[i],
                    model_predictions[j]
                )
                mcnemar_result['model1'] = model_names[i]
                mcnemar_result['model2'] = model_names[j]
                mcnemar_results.append(mcnemar_result)

    return {
        'matthews_coefficients': matthews_coefficients,
        'mcnemar_results': mcnemar_results,
        'sample_size': len(y_true_real),
        'real_data': True
    }

def generate_interpretation_for_professor(mcnemar_analysis, validation_data):
    """
    Genera interpretación concisa para el profesor
    """
    if not mcnemar_analysis:
        return "No hay datos para interpretar."

    # Análisis básico
    sample_size = mcnemar_analysis['sample_size']
    matthews_coefficients = mcnemar_analysis['matthews_coefficients']
    mcnemar_results = mcnemar_analysis['mcnemar_results']

    # Encontrar mejor modelo por MCC
    best_mcc_model = max(matthews_coefficients, key=lambda x: x['mcc'])

    # Encontrar mejor modelo por precisión
    y_true = validation_data['y_true']
    model_predictions = validation_data['predictions']
    model_names = validation_data['model_names']

    accuracies = []
    for i, (model_name, predictions) in enumerate(zip(model_names, model_predictions)):
        accuracy = np.mean(y_true == predictions)
        accuracies.append({'model': model_name, 'accuracy': accuracy})

    best_accuracy_model = max(accuracies, key=lambda x: x['accuracy'])

    # Contar diferencias significativas
    significant_differences = len([r for r in mcnemar_results if r['p_value'] < 0.05])

    # Generar interpretación
    interpretation = f"""
**INTERPRETACIÓN PARA PRESENTACIÓN ACADÉMICA**

**Dataset de Validación:** {sample_size} imágenes reales de hojas de vid

**Modelo Recomendado:** {best_accuracy_model['model']} (Precisión: {best_accuracy_model['accuracy']:.1%})

**Análisis Estadístico:**
• **Coeficiente de Matthews (MCC):** {best_mcc_model['mcc']:.3f} - {best_mcc_model['interpretation']}
• **Pruebas de McNemar:** {significant_differences} de {len(mcnemar_results)} comparaciones muestran diferencias significativas (p < 0.05)

**Conclusión Científica:**
"""

    if significant_differences > 0:
        interpretation += f"Existen diferencias estadísticamente significativas entre algunos modelos, validando la necesidad de selección cuidadosa del algoritmo. {best_accuracy_model['model']} muestra el mejor rendimiento general."
    else:
        interpretation += f"No se encontraron diferencias estadísticamente significativas entre modelos (p ≥ 0.05), indicando rendimiento equivalente. Cualquier modelo es válido para implementación clínica."

    if best_mcc_model['mcc'] == 0:
        interpretation += f"\n\n**Nota Metodológica:** MCC = 0 indica dataset homogéneo (una clase predominante), típico en validaciones clínicas enfocadas."

    return interpretation

# ======= FIN FUNCIONES PARA VALIDACIÓN =======

# Función para generar recomendaciones
def get_treatment_recommendations(disease, language='es'):
    """Obtiene recomendaciones de tratamiento según la enfermedad e idioma"""
    
    # Recomendaciones en español
    recommendations_es = {
        "Black_rot": {
            "titulo": "🔴 Podredumbre Negra Detectada",
            "gravedad": "Alta",
            "tratamiento": [
                "Aplicar fungicidas protectores (Mancozeb, Captan)",
                "Eliminar y destruir todas las partes infectadas",
                "Mejorar la circulación de aire en el viñedo",
                "Evitar el riego por aspersión"
            ],
            "prevencion": [
                "Podar adecuadamente para mejorar ventilación",
                "Aplicar fungicidas preventivos antes de la floración",
                "Eliminar restos de poda y hojas caídas"
            ]
        },
        "Esca": {
            "titulo": "🟤 Esca (Sarampión Negro) Detectada",
            "gravedad": "Muy Alta",
            "tratamiento": [
                "No existe cura directa - enfoque en prevención",
                "Podar las partes afectadas con herramientas desinfectadas",
                "Aplicar pasta cicatrizante en cortes de poda",
                "Considerar reemplazo de plantas severamente afectadas"
            ],
            "prevencion": [
                "Evitar podas tardías y en días húmedos",
                "Desinfectar herramientas entre plantas",
                "Proteger heridas de poda inmediatamente"
            ]
        },
        "Healthy": {
            "titulo": "✅ Planta Sana",
            "gravedad": "Ninguna",
            "tratamiento": [
                "No se requiere tratamiento",
                "Mantener las prácticas actuales de manejo"
            ],
            "prevencion": [
                "Continuar monitoreo regular",
                "Mantener programa preventivo de fungicidas",
                "Asegurar nutrición balanceada",
                "Mantener buen drenaje del suelo"
            ]
        },
        "Leaf_blight": {
            "titulo": "🟡 Tizón de la Hoja Detectado",
            "gravedad": "Moderada",
            "tratamiento": [
                "Aplicar fungicidas sistémicos (Azoxistrobina, Tebuconazol)",
                "Remover hojas infectadas",
                "Mejorar el drenaje del suelo",
                "Reducir la densidad del follaje"
            ],
            "prevencion": [
                "Evitar el exceso de nitrógeno",
                "Mantener el follaje seco",
                "Aplicar fungicidas preventivos en épocas húmedas"
            ]
        }
    }
    
    # Recomendaciones en inglés
    recommendations_en = {
        "Black_rot": {
            "titulo": "🔴 Black Rot Detected",
            "gravedad": "High",
            "tratamiento": [
                "Apply protective fungicides (Mancozeb, Captan)",
                "Remove and destroy all infected parts",
                "Improve air circulation in the vineyard",
                "Avoid overhead irrigation"
            ],
            "prevencion": [
                "Prune properly to improve ventilation",
                "Apply preventive fungicides before flowering",
                "Remove pruning debris and fallen leaves"
            ]
        },
        "Esca": {
            "titulo": "🟤 Esca (Black Measles) Detected",
            "gravedad": "Very High",
            "tratamiento": [
                "No direct cure - focus on prevention",
                "Prune affected parts with disinfected tools",
                "Apply healing paste on pruning cuts",
                "Consider replacement of severely affected plants"
            ],
            "prevencion": [
                "Avoid late pruning on humid days",
                "Disinfect tools between plants",
                "Protect pruning wounds immediately"
            ]
        },
        "Healthy": {
            "titulo": "✅ Healthy Plant",
            "gravedad": "None",
            "tratamiento": [
                "No treatment required",
                "Maintain current management practices"
            ],
            "prevencion": [
                "Continue regular monitoring",
                "Maintain preventive fungicide program",
                "Ensure balanced nutrition",
                "Maintain good soil drainage"
            ]
        },
        "Leaf_blight": {
            "titulo": "🟡 Leaf Blight Detected",
            "gravedad": "Moderate",
            "tratamiento": [
                "Apply systemic fungicides (Azoxystrobin, Tebuconazole)",
                "Remove infected leaves",
                "Improve soil drainage",
                "Reduce foliage density"
            ],
            "prevencion": [
                "Avoid excess nitrogen",
                "Keep foliage dry",
                "Apply preventive fungicides in humid periods"
            ]
        }
    }
    
    # Recomendaciones en portugués
    recommendations_pt = {
        "Black_rot": {
            "titulo": "🔴 Podridão Negra Detectada",
            "gravedad": "Alta",
            "tratamiento": [
                "Aplicar fungicidas protetores (Mancozeb, Captan)",
                "Eliminar e destruir todas as partes infectadas",
                "Melhorar a circulação de ar no vinhedo",
                "Evitar irrigação por aspersão"
            ],
            "prevencion": [
                "Podar adequadamente para melhorar ventilação",
                "Aplicar fungicidas preventivos antes da floração",
                "Eliminar restos de poda e folhas caídas"
            ]
        },
        "Esca": {
            "titulo": "🟤 Esca (Sarampo Negro) Detectada",
            "gravedad": "Muito Alta",
            "tratamiento": [
                "Não existe cura direta - foco na prevenção",
                "Podar partes afetadas com ferramentas desinfetadas",
                "Aplicar pasta cicatrizante em cortes de poda",
                "Considerar substituição de plantas severamente afetadas"
            ],
            "prevencion": [
                "Evitar podas tardias em dias úmidos",
                "Desinfetar ferramentas entre plantas",
                "Proteger feridas de poda imediatamente"
            ]
        },
        "Healthy": {
            "titulo": "✅ Planta Saudável",
            "gravedad": "Nenhuma",
            "tratamiento": [
                "Não é necessário tratamento",
                "Manter as práticas atuais de manejo"
            ],
            "prevencion": [
                "Continuar monitoramento regular",
                "Manter programa preventivo de fungicidas",
                "Assegurar nutrição balanceada",
                "Manter boa drenagem do solo"
            ]
        },
        "Leaf_blight": {
            "titulo": "🟡 Queima das Folhas Detectada",
            "gravedad": "Moderada",
            "tratamiento": [
                "Aplicar fungicidas sistêmicos (Azoxistrobina, Tebuconazol)",
                "Remover folhas infectadas",
                "Melhorar a drenagem do solo",
                "Reduzir a densidade da folhagem"
            ],
            "prevencion": [
                "Evitar excesso de nitrogênio",
                "Manter folhagem seca",
                "Aplicar fungicidas preventivos em períodos úmidos"
            ]
        }
    }
    
    # Seleccionar recomendaciones según idioma
    if language == 'en':
        recommendations = recommendations_en
    elif language == 'pt':
        recommendations = recommendations_pt
    else:
        recommendations = recommendations_es
    
    return recommendations.get(disease, {})

# ======= FUNCIÓN PDF MEJORADA (SIN ANÁLISIS ESTADÍSTICO) =======
def generate_diagnosis_pdf(image, results, consensus_disease):
    """Genera un reporte PDF del diagnóstico sin análisis estadístico"""
    
    # Obtener idioma actual
    current_language = st.session_state.language
    
    # Obtener recomendaciones en el idioma actual
    recommendations = get_treatment_recommendations(consensus_disease, current_language)

    # Datos de entrenamiento de los 5 modelos
    training_data = {
        "M1 - SVM": {"tipo": "Clásico", "description": "SVM kernel RBF", "features": "Color + LBP + Stats", "epochs": "N/A", "time": "N/A", "accuracy": "N/A", "val_accuracy": "N/A"},
        "M2 - Random Forest": {"tipo": "Clásico", "description": "RF 200 árboles", "features": "Color + LBP + Stats", "epochs": "N/A", "time": "N/A", "accuracy": "N/A", "val_accuracy": "N/A"},
        "M3 - KNN": {"tipo": "Clásico", "description": "KNN k=5", "features": "Color + LBP + Stats", "epochs": "N/A", "time": "N/A", "accuracy": "N/A", "val_accuracy": "N/A"},
        "H1 - CNN + SVM": {"tipo": "Híbrido", "description": "CNN extractor + SVM", "features": "Deep CNN (256 dims)", "epochs": "N/A", "time": "N/A", "accuracy": "N/A", "val_accuracy": "N/A"},
        "H2 - Transfer + RF": {"tipo": "Híbrido", "description": "MobileNetV2 + RF", "features": "ImageNet (1280 dims)", "epochs": "N/A", "time": "N/A", "accuracy": "N/A", "val_accuracy": "N/A"},
    }

    # Crear archivo temporal para el PDF
    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
        pdf_filename = tmp_file.name

    try:
        with PdfPages(pdf_filename) as pdf:

            # ====================== PÁGINA 1: PORTADA ======================
            fig = plt.figure(figsize=(8.27, 11.69))
            fig.patch.set_facecolor('white')

            # Título principal
            fig.text(0.5, 0.9, get_text('title', current_language), fontsize=24, fontweight='bold',
                     ha='center', color='#2E8B57')
            
            # Subtítulo según idioma
            if current_language == 'en':
                subtitle = 'Vineyard Disease Diagnosis Report'
            elif current_language == 'pt':
                subtitle = 'Relatório de Diagnóstico de Doenças em Vinhedos'
            else:  # español
                subtitle = 'Reporte de Diagnóstico de Enfermedades en Viñedos'
                
            fig.text(0.5, 0.85, subtitle, fontsize=14, ha='center', color='#333333')

            # Información del reporte
            date_label = 'Date:' if current_language == 'en' else 'Data:' if current_language == 'pt' else 'Fecha:'
            models_label = 'Models used:' if current_language == 'en' else 'Modelos utilizados:' if current_language == 'pt' else 'Modelos utilizados:'
            
            fig.text(0.1, 0.75, f'{date_label} {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', fontsize=11)
            fig.text(0.1, 0.72, f'{models_label} {len(results)}', fontsize=11)

            # Diagnóstico principal
            predictions = [r['predicted_class'] for r in results]
            consensus = max(set(predictions), key=predictions.count)
            consensus_count = predictions.count(consensus)
            consensus_confidence = np.mean([r['confidence'] for r in results if r['predicted_class'] == consensus])

            # Etiquetas traducidas
            main_diagnosis_label = 'MAIN DIAGNOSIS' if current_language == 'en' else 'DIAGNÓSTICO PRINCIPAL' if current_language == 'pt' else 'DIAGNÓSTICO PRINCIPAL'
            disease_label = 'Disease:' if current_language == 'en' else 'Doença:' if current_language == 'pt' else 'Enfermedad:'
            confidence_label = 'Confidence:' if current_language == 'en' else 'Confiança:' if current_language == 'pt' else 'Confianza:'
            consensus_label = 'Consensus:' if current_language == 'en' else 'Consenso:' if current_language == 'pt' else 'Consenso:'

            fig.text(0.1, 0.6, main_diagnosis_label, fontsize=16, fontweight='bold', color='#2E8B57')
            fig.text(0.1, 0.55, f'{disease_label} {get_disease_names(current_language)[consensus]}', fontsize=12)
            fig.text(0.1, 0.52, f'{confidence_label} {consensus_confidence:.1%}', fontsize=12)
            fig.text(0.1, 0.49, f'{consensus_label} {consensus_count}/{len(results)} modelos', fontsize=12)

            # Recomendaciones clave
            if recommendations:
                key_recommendations_label = 'KEY RECOMMENDATIONS' if current_language == 'en' else 'RECOMENDAÇÕES PRINCIPAIS' if current_language == 'pt' else 'RECOMENDACIONES CLAVE'
                severity_label = 'Severity:' if current_language == 'en' else 'Gravidade:' if current_language == 'pt' else 'Gravedad:'
                action_label = 'Action:' if current_language == 'en' else 'Ação:' if current_language == 'pt' else 'Acción:'
                
                fig.text(0.1, 0.4, key_recommendations_label, fontsize=14, fontweight='bold', color='#2E8B57')
                fig.text(0.1, 0.35, f'{severity_label} {recommendations.get("gravedad", "N/A")}', fontsize=11)
                action = recommendations.get('tratamiento', ['N/A'])[0] if recommendations.get('tratamiento') else 'N/A'
                if len(action) > 60:
                    action = action[:60] + "..."
                fig.text(0.1, 0.32, f'{action_label} {action}', fontsize=10)

            plt.axis('off')
            pdf.savefig(fig, bbox_inches='tight')
            plt.close()

            # ====================== PÁGINA 2: RESULTADOS DETALLADOS ======================
            fig, ((ax1, ax2), (ax3, ax4)) = plt.subplots(2, 2, figsize=(8.27, 11.69))
            
            # Título traducido
            detailed_analysis_title = 'Detailed Model Analysis' if current_language == 'en' else 'Análise Detalhada de Modelos' if current_language == 'pt' else 'Análisis Detallado de Modelos'
            fig.suptitle(detailed_analysis_title, fontsize=16, fontweight='bold')

            # Gráfico 1: Confianza por modelo
            model_names = [r['model_name'] for r in results]
            confidences = [r['confidence'] for r in results]
            colors = ['#3498db', '#2ecc71', '#e74c3c', '#f39c12']

            bars1 = ax1.bar(range(len(model_names)), confidences, color=colors)
            
            # Títulos traducidos
            confidence_title = 'Confidence by Model' if current_language == 'en' else 'Confiança por Modelo' if current_language == 'pt' else 'Confianza por Modelo'
            confidence_ylabel = 'Confidence' if current_language == 'en' else 'Confiança' if current_language == 'pt' else 'Confianza'
            
            ax1.set_title(confidence_title)
            ax1.set_ylabel(confidence_ylabel)
            ax1.set_xticks(range(len(model_names)))
            ax1.set_xticklabels([name.replace(' ', '\n') for name in model_names], fontsize=9)

            for bar, conf in zip(bars1, confidences):
                height = bar.get_height()
                ax1.text(bar.get_x() + bar.get_width()/2., height + 0.01,
                         f'{conf:.1%}', ha='center', va='bottom', fontweight='bold')

            # Gráfico 2: Tiempo de inferencia
            inference_times = [r['inference_time'] for r in results]
            bars2 = ax2.bar(range(len(model_names)), inference_times, color=colors)
            
            # Títulos traducidos
            inference_time_title = 'Inference Time (ms)' if current_language == 'en' else 'Tempo de Inferência (ms)' if current_language == 'pt' else 'Tiempo de Inferencia (ms)'
            time_ylabel = 'Time (ms)' if current_language == 'en' else 'Tempo (ms)' if current_language == 'pt' else 'Tiempo (ms)'
            
            ax2.set_title(inference_time_title)
            ax2.set_ylabel(time_ylabel)
            ax2.set_xticks(range(len(model_names)))
            ax2.set_xticklabels([name.replace(' ', '\n') for name in model_names], fontsize=9)

            for bar, time in zip(bars2, inference_times):
                height = bar.get_height()
                ax2.text(bar.get_x() + bar.get_width()/2., height + 1,
                         f'{time:.0f}', ha='center', va='bottom', fontweight='bold')

            # Gráfico 3: Distribución de probabilidades
            best_result = max(results, key=lambda x: x['confidence'])
            all_probs = best_result['all_predictions']
            
            # Usar nombres de enfermedades traducidos
            disease_names_translated = [get_disease_names(current_language)[cls] for cls in DISEASE_CLASSES]

            wedges, texts, autotexts = ax3.pie(all_probs, labels=disease_names_translated,
                                               autopct='%1.1f%%', startangle=90,
                                               colors=['#FFB6C1', '#98FB98', '#87CEEB', '#DDA0DD'])
            
            # Título traducido
            probabilities_title = f'Probabilities\n({best_result["model_name"]})' if current_language == 'en' else f'Probabilidades\n({best_result["model_name"]})' if current_language == 'pt' else f'Probabilidades\n({best_result["model_name"]})'
            ax3.set_title(probabilities_title)

            # Gráfico 4: Consenso entre modelos
            consensus_data = {}
            for pred in predictions:
                consensus_data[pred] = consensus_data.get(pred, 0) + 1

            labels = [get_disease_names(current_language)[k] for k in consensus_data.keys()]
            values = list(consensus_data.values())

            bars4 = ax4.bar(range(len(labels)), values, color=['#FF6B6B', '#4ECDC4', '#45B7D1', '#96CEB4'])
            
            # Títulos traducidos
            consensus_title = 'Model Consensus' if current_language == 'en' else 'Consenso entre Modelos' if current_language == 'pt' else 'Consenso entre Modelos'
            models_ylabel = 'Number of Models' if current_language == 'en' else 'Número de Modelos' if current_language == 'pt' else 'Número de Modelos'
            
            ax4.set_title(consensus_title)
            ax4.set_ylabel(models_ylabel)
            ax4.set_xticks(range(len(labels)))
            ax4.set_xticklabels([label.replace(' ', '\n') for label in labels], fontsize=8)

            for bar, val in zip(bars4, values):
                height = bar.get_height()
                ax4.text(bar.get_x() + bar.get_width()/2., height + 0.05,
                         str(val), ha='center', va='bottom', fontweight='bold')

            plt.tight_layout()
            pdf.savefig(fig, bbox_inches='tight')
            plt.close()

            # ====================== PÁGINA 3: MATRIZ DE CONFUSIÓN Y ENTRENAMIENTO ======================
            fig = plt.figure(figsize=(8.27, 11.69))

            # Título traducido
            confusion_training_title = 'Confusion Matrix and Training Data' if current_language == 'en' else 'Matriz de Confusão e Dados de Treinamento' if current_language == 'pt' else 'Matriz de Confusión y Datos de Entrenamiento'
            fig.text(0.5, 0.95, confusion_training_title, fontsize=16, fontweight='bold', ha='center')

            ax_matrix = fig.add_subplot(2, 1, 1)

            cm_csv_map = {
                "M1 - SVM": "confusion_m1_svm.csv",
                "M2 - Random Forest": "confusion_m2_random_forest.csv",
                "M3 - KNN": "confusion_m3_knn.csv",
                "H1 - CNN + SVM": "confusion_h1_cnn_svm.csv",
                "H2 - Transfer + RF": "confusion_h2_transfer_rf.csv",  # nombre debe coincidir EXACTO con MODELS_CONFIG
            }
            cm_filename = cm_csv_map.get(best_result["model_name"])
            cm_path = Path("reports/modelos") / cm_filename if cm_filename else None
            cm_is_real = cm_path is not None and cm_path.is_file()
            if cm_is_real:
                confusion_matrix_data = pd.read_csv(cm_path, index_col=0).values
            else:
                confusion_matrix_data = np.identity(len(DISEASE_CLASSES), dtype=int)

            im = ax_matrix.imshow(confusion_matrix_data, interpolation='nearest', cmap='Blues')
            
            confusion_matrix_title = f'Confusion Matrix - {best_result["model_name"]}' if current_language == 'en' else f'Matriz de Confusão - {best_result["model_name"]}' if current_language == 'pt' else f'Matriz de Confusión - {best_result["model_name"]}'
            if not cm_is_real:
                no_data_note = ' (sin datos de test disponibles)' if current_language != 'en' else ' (test data not available)'
                confusion_matrix_title += no_data_note
            ax_matrix.set_title(confusion_matrix_title, fontweight='bold', pad=20, fontsize=11)

            class_names_translated = [get_disease_names(current_language)[cls] for cls in DISEASE_CLASSES]
            ax_matrix.set_xticks(range(len(class_names_translated)))
            ax_matrix.set_yticks(range(len(class_names_translated)))
            ax_matrix.set_xticklabels(class_names_translated)
            ax_matrix.set_yticklabels(class_names_translated)
            
            prediction_label = 'Prediction' if current_language == 'en' else 'Predição' if current_language == 'pt' else 'Predicción'
            actual_label = 'Actual' if current_language == 'en' else 'Real' if current_language == 'pt' else 'Real'
            
            ax_matrix.set_xlabel(prediction_label, fontweight='bold')
            ax_matrix.set_ylabel(actual_label, fontweight='bold')

            vmax = confusion_matrix_data.max()
            for i in range(len(class_names_translated)):
                for j in range(len(class_names_translated)):
                    text = ax_matrix.text(j, i, confusion_matrix_data[i, j],
                                          ha="center", va="center",
                                          color="white" if confusion_matrix_data[i, j] > vmax * 0.5 else "black",
                                          fontweight='bold')

            # Tabla de entrenamiento
            ax_table = fig.add_subplot(2, 1, 2)
            ax_table.axis('tight')
            ax_table.axis('off')

            # Crear tabla de información de entrenamiento
            table_data = []
            headers = ['Modelo', 'Epochs', 'Tiempo', 'Precisión', 'Val. Precisión', 'Inferencia']

            for result in results:
                model_name = result['model_name']
                train_info = training_data.get(model_name, {"epochs": "N/A", "time": "N/A",
                                                            "accuracy": "N/A", "val_accuracy": "N/A"})
                table_data.append([
                    model_name,
                    train_info['epochs'],
                    train_info['time'],
                    train_info['accuracy'],
                    train_info['val_accuracy'],
                    f"{result['inference_time']:.0f} ms"
                ])

            table = ax_table.table(cellText=table_data,
                                   colLabels=headers,
                                   cellLoc='center',
                                   loc='center')

            table.auto_set_font_size(False)
            table.set_fontsize(10)
            table.scale(1, 2)

            # Colorear encabezados
            for i in range(len(headers)):
                table[(0, i)].set_facecolor('#2E8B57')
                table[(0, i)].set_text_props(weight='bold', color='white')

            ax_table.set_title('Información de Entrenamiento y Rendimiento',
                               fontweight='bold', fontsize=14, pad=20)

            pdf.savefig(fig, bbox_inches='tight')
            plt.close()

            # ====================== PÁGINA 4: RECOMENDACIONES ======================
            fig = plt.figure(figsize=(8.27, 11.69))
            treatment_recommendations_title = 'Treatment Recommendations' if current_language == 'en' else 'Recomendações de Tratamento' if current_language == 'pt' else 'Recomendaciones de Tratamiento'
            fig.text(0.5, 0.95, treatment_recommendations_title, fontsize=16, fontweight='bold', ha='center')

            if recommendations:
                fig.text(0.1, 0.85, recommendations.get('titulo', ''), fontsize=14, fontweight='bold', color='#B22222')
                
                severity_label = 'Severity:' if current_language == 'en' else 'Gravidade:' if current_language == 'pt' else 'Gravedad:'
                fig.text(0.1, 0.8, f"{severity_label} {recommendations.get('gravedad', 'N/A')}", fontsize=12, fontweight='bold')

                # Tratamientos
                recommended_treatments_label = 'RECOMMENDED TREATMENTS:' if current_language == 'en' else 'TRATAMENTOS RECOMENDADOS:' if current_language == 'pt' else 'TRATAMIENTOS RECOMENDADOS:'
                fig.text(0.1, 0.7, recommended_treatments_label, fontsize=12, fontweight='bold')
                y_pos = 0.65
                for i, item in enumerate(recommendations.get('tratamiento', []), 1):
                    fig.text(0.1, y_pos, f"{i}. {item}", fontsize=10)
                    y_pos -= 0.04

                # Prevención
                preventive_measures_label = 'PREVENTIVE MEASURES:' if current_language == 'en' else 'MEDIDAS PREVENTIVAS:' if current_language == 'pt' else 'MEDIDAS PREVENTIVAS:'
                fig.text(0.1, 0.4, preventive_measures_label, fontsize=12, fontweight='bold')
                y_pos = 0.35
                for i, item in enumerate(recommendations.get('prevencion', []), 1):
                    fig.text(0.1, y_pos, f"{i}. {item}", fontsize=10)
                    y_pos -= 0.04

            # Nota
            note_text = 'Note: Consult with a specialist before applying treatments.' if current_language == 'en' else 'Nota: Consulte um especialista antes de aplicar tratamentos.' if current_language == 'pt' else 'Nota: Consulte con un especialista antes de aplicar tratamientos.'
            fig.text(0.1, 0.1, note_text, fontsize=10, style='italic')

            plt.axis('off')
            pdf.savefig(fig, bbox_inches='tight')
            plt.close()

        # Leer el archivo PDF generado
        with open(pdf_filename, 'rb') as f:
            pdf_bytes = f.read()

        return pdf_bytes

    finally:
        # Limpiar archivo temporal
        if os.path.exists(pdf_filename):
            os.unlink(pdf_filename)

# ======= FUNCIÓN WORD =======
def generate_diagnosis_docx(image, results, consensus_disease):
    """Genera un reporte Word del diagnóstico"""
    current_language = st.session_state.language
    recommendations = get_treatment_recommendations(consensus_disease, current_language)
    
    # Crear documento
    doc = Document()
    
    # Título
    title = doc.add_heading(get_text('title', current_language), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtítulo
    if current_language == 'en':
        subtitle = 'Vineyard Disease Diagnosis Report'
    elif current_language == 'pt':
        subtitle = 'Relatório de Diagnóstico de Doenças em Vinhedos'
    else:
        subtitle = 'Reporte de Diagnóstico de Enfermedades en Viñedos'
    doc.add_heading(subtitle, level=1)
    
    # Información del reporte
    date_label = 'Date:' if current_language == 'en' else 'Data:' if current_language == 'pt' else 'Fecha:'
    doc.add_paragraph(f'{date_label} {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    
    doc.add_heading('Diagnóstico Principal', level=2)
    doc.add_paragraph(f'{get_disease_names(current_language)[consensus_disease]}')
    
    # Resultados por modelo
    doc.add_heading('Resultados por Modelo', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Modelo'
    hdr_cells[1].text = 'Diagnóstico'
    hdr_cells[2].text = 'Confianza'
    hdr_cells[3].text = 'Tiempo (ms)'
    
    for result in results:
        row_cells = table.add_row().cells
        row_cells[0].text = result['model_name']
        row_cells[1].text = get_disease_names(current_language)[result['predicted_class']]
        row_cells[2].text = f'{result["confidence"]:.1%}'
        row_cells[3].text = f'{result["inference_time"]:.0f}'
    
    # Recomendaciones
    if recommendations:
        doc.add_heading('Recomendaciones', level=2)
        doc.add_heading('Gravedad', level=3)
        doc.add_paragraph(recommendations.get('gravedad', 'N/A'))
        doc.add_heading('Tratamiento', level=3)
        for item in recommendations.get('tratamiento', []):
            doc.add_paragraph(item, style='List Bullet')
        doc.add_heading('Prevención', level=3)
        for item in recommendations.get('prevencion', []):
            doc.add_paragraph(item, style='List Bullet')
    
    # Guardar documento
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
        doc.save(tmp_file.name)
        docx_filename = tmp_file.name
    
    try:
        with open(docx_filename, 'rb') as f:
            docx_bytes = f.read()
        return docx_bytes
    finally:
        if os.path.exists(docx_filename):
            os.unlink(docx_filename)

# ======= FUNCIÓN EXCEL =======
def generate_diagnosis_xlsx(image, results, consensus_disease):
    """Genera un reporte Excel del diagnóstico"""
    current_language = st.session_state.language
    recommendations = get_treatment_recommendations(consensus_disease, current_language)
    
    # Crear libro de Excel
    wb = Workbook()
    ws = wb.active
    ws.title = "Reporte"
    
    # Estilos
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="2E8B57", end_color="2E8B57", fill_type="solid")
    center_alignment = Alignment(horizontal="center", vertical="center")
    
    # Título
    ws['A1'] = get_text('title', current_language)
    ws.merge_cells('A1:D1')
    cell = ws['A1']
    cell.font = Font(bold=True, size=16, color="2E8B57")
    cell.alignment = center_alignment
    
    # Fecha
    date_label = 'Date:' if current_language == 'en' else 'Data:' if current_language == 'pt' else 'Fecha:'
    ws['A2'] = f'{date_label} {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}'
    ws.merge_cells('A2:D2')
    
    # Diagnóstico principal
    ws['A4'] = 'Diagnóstico Principal:'
    ws['A4'].font = header_font
    ws['A4'].fill = header_fill
    ws.merge_cells('A4:B4')
    ws['C4'] = get_disease_names(current_language)[consensus_disease]
    ws.merge_cells('C4:D4')
    
    # Resultados por modelo
    ws['A6'] = 'Resultados por Modelo'
    ws['A6'].font = header_font
    ws['A6'].fill = header_fill
    ws.merge_cells('A6:D6')
    
    headers = ['Modelo', 'Diagnóstico', 'Confianza', 'Tiempo (ms)']
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=7, column=col_num, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center_alignment
    
    for row_num, result in enumerate(results, 8):
        ws.cell(row=row_num, column=1, value=result['model_name'])
        ws.cell(row=row_num, column=2, value=get_disease_names(current_language)[result['predicted_class']])
        ws.cell(row=row_num, column=3, value=f'{result["confidence"]:.1%}')
        ws.cell(row=row_num, column=4, value=f'{result["inference_time"]:.0f}')
    
    # Ajustar ancho de columnas
    for col in ws.columns:
        max_length = 0
        column = col[0].column_letter
        for cell in col:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = (max_length + 2)
        ws.column_dimensions[column].width = adjusted_width
    
    # Guardar Excel
    with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp_file:
        wb.save(tmp_file.name)
        xlsx_filename = tmp_file.name
    
    try:
        with open(xlsx_filename, 'rb') as f:
            xlsx_bytes = f.read()
        return xlsx_bytes
    finally:
        if os.path.exists(xlsx_filename):
            os.unlink(xlsx_filename)

# INTERFAZ PRINCIPAL
def main():
    # Título y descripción
    st.title(get_text('title', st.session_state.language))
    st.markdown(f"**{get_text('subtitle', st.session_state.language)}**")
    st.markdown(f"*{get_text('subtitle_analysis', st.session_state.language)}*")

    # Sidebar
    with st.sidebar:
        st.header(get_text('config_title', st.session_state.language))

        # Cargar modelos si no están cargados
        if not st.session_state.models_loaded:
            if st.button(get_text('load_models', st.session_state.language), type="primary"):
                loading_msg = 'Cargando modelos...' if st.session_state.language == 'es' else 'Loading models...' if st.session_state.language == 'en' else 'Carregando modelos...'
                with st.spinner(loading_msg):
                    st.session_state.models = load_models()
                    if st.session_state.models:
                        st.session_state.models_loaded = True
                        success_msg = "✅ Modelos cargados exitosamente!" if st.session_state.language == 'es' else "✅ Models loaded successfully!" if st.session_state.language == 'en' else "✅ Modelos carregados com sucesso!"
                        st.success(success_msg)
                    else:
                        error_msg = "❌ No se pudieron cargar los modelos" if st.session_state.language == 'es' else "❌ Could not load models" if st.session_state.language == 'en' else "❌ Não foi possível carregar os modelos"
                        st.error(error_msg)
        else:
            st.success(get_text('models_ready', st.session_state.language))

        # ── Panel de Sistema ──────────────────────────────────────────────
        lang = st.session_state.language
        t = lambda es, en, pt: es if lang == 'es' else en if lang == 'en' else pt

        status_class = "loaded" if st.session_state.models_loaded else "not-loaded"
        status_text = t("Cargados", "Loaded", "Carregados") if st.session_state.models_loaded else t("No cargados", "Not loaded", "Não carregados")

        models_list_html = "".join([f"<li>{name}</li>" for name in MODEL_PATHS.keys()])
        classes_list_html = "".join([f"<li>{get_disease_names(st.session_state.language)[cls]}</li>" for cls in DISEASE_CLASSES])

        # ── Mejor modelo ──────────────────────────────────────────────────
        best_model_card = ""
        if st.session_state.best_model_name:
            name = st.session_state.best_model_name.split(":")[1].strip() if ":" in st.session_state.best_model_name else st.session_state.best_model_name
            best_model_card = f"""
            <div style="background:linear-gradient(135deg,#4a148c,#2e7d32);border-radius:12px;padding:14px;margin:12px 0;text-align:center;">
                <div style="color:#ffd700;font-size:0.75rem;font-weight:700;text-transform:uppercase;letter-spacing:1px;">
                    {t("🏆 MEJOR MODELO","🏆 BEST MODEL","🏆 MELHOR MODELO")}
                </div>
                <div style="color:white;font-size:1.1rem;font-weight:800;margin-top:4px;">{name}</div>
            </div>"""

        # ── Ranking rápido ─────────────────────────────────────────────────
        ranking_html = ""
        if st.session_state.ranking_data:
            top3 = st.session_state.ranking_data[:3]
            items = []
            for i, m in enumerate(top3):
                rank_icon = ["🥇", "🥈", "🥉"][i]
                raw = m.get("accuracy", m.get("puntaje", ""))
                if isinstance(raw, (float, int)):
                    score_str = f"{float(raw):.4f}"
                else:
                    score_str = str(raw)
                items.append(f"<div style='display:flex;justify-content:space-between;padding:2px 0;font-size:0.8rem;'><span>{rank_icon} {m.get('modelo','')}</span><span style='font-weight:600;color:#4a148c;'>{score_str}</span></div>")
            ranking_html = f"""
            <div style="background:#f9f9fb;border-radius:10px;padding:10px;margin:8px 0;border:1px solid #edeef2;">
                <div style="font-size:0.8rem;font-weight:700;color:#4a148c;margin-bottom:6px;">
                    {t("📊 Ranking","📊 Ranking","📊 Ranking")}
                </div>
                {''.join(items)}
            </div>"""

        # ── Pipeline status ────────────────────────────────────────────────
        pipeline = check_pipeline_status()
        pipe_items = []
        for key, p in pipeline.items():
            icon = "✅" if p["done"] else "⬜"
            pipe_items.append(f"<div style='display:flex;align-items:center;gap:6px;padding:1px 0;font-size:0.78rem;color:#424242;'><span>{icon}</span><span>{p['label']}</span></div>")

        pipeline_html = f"""
        <div style="background:#f9f9fb;border-radius:10px;padding:10px;margin:8px 0;border:1px solid #edeef2;">
            <div style="font-size:0.8rem;font-weight:700;color:#4a148c;margin-bottom:6px;">
                {t("⚙️ Pipeline de experimentos","⚙️ Experiment pipeline","⚙️ Pipeline de experimentos")}
            </div>
            {''.join(pipe_items)}
        </div>"""

        # ── Info block ─────────────────────────────────────────────────────
        info_block_html = f"""
        <div class="sidebar-info-block">
            <h4 class="sidebar-info-title">{t("📊 Info. del Sistema","📊 System Info","📊 Info. do Sistema")}</h4>
            <p class="sidebar-info-item"><span class="sidebar-info-label">{t("Estado:","Status:","Estado:")}</span> <span class="status-badge {status_class}">{status_text}</span></p>
            <p class="sidebar-info-item"><span class="sidebar-info-label">{t("Modelos:","Models:","Modelos:")}</span>
                <ul style="margin:2px 0;padding-left:20px;font-size:0.8rem;color:#424242;">{models_list_html}</ul>
            </p>
            <p class="sidebar-info-item"><span class="sidebar-info-label">{t("Clases:","Classes:","Classes:")}</span>
                <ul style="margin:2px 0;padding-left:20px;font-size:0.8rem;color:#424242;">{classes_list_html}</ul>
            </p>
        </div>
        """
        st.markdown(info_block_html, unsafe_allow_html=True)

        # ── Best model card ────────────────────────────────────────────────
        if best_model_card:
            st.markdown(best_model_card, unsafe_allow_html=True)

        # ── Ranking ────────────────────────────────────────────────────────
        if ranking_html:
            st.markdown(ranking_html, unsafe_allow_html=True)

        # ── Pipeline ───────────────────────────────────────────────────────
        st.markdown(pipeline_html, unsafe_allow_html=True)

        # Información
        st.markdown("---")
        st.subheader(get_text('info_title', st.session_state.language))
        info_desc = get_text('info_description', st.session_state.language).replace('•', '▪').replace('\n', '<br>')
        st.markdown(f"""
        <div style="font-size: 0.85rem; color: #555; line-height: 1.4;">
            {info_desc}
        </div>
        """, unsafe_allow_html=True)

    # Contenido principal cuando los modelos no están cargados
    if not st.session_state.models_loaded:
        # Hero Section
        st.markdown(f"""
        <div class="hero-container">
            <h1 class="hero-title">🍇 VineGuard AI</h1>
            <h3 class="hero-subtitle">{"Diagnóstico inteligente de enfermedades foliares en vid mediante Deep Learning" if st.session_state.language == 'es' else "Intelligent grapevine leaf disease diagnosis using Deep Learning" if st.session_state.language == 'en' else "Diagnóstico inteligente de doenças foliares em videira usando Deep Learning"}</h3>
            <p class="hero-description">
                {"Esta plataforma utiliza redes neuronales convolucionales (CNN) avanzadas para identificar de forma precisa y rápida patologías comunes en hojas de vid. Cargue una imagen, compare predicciones de múltiples modelos en tiempo real y genere un informe completo con recomendaciones de tratamiento." if st.session_state.language == 'es' else "This platform utilizes advanced convolutional neural networks (CNN) to accurately and rapidly identify common grapevine leaf pathologies. Upload an image, compare predictions from multiple models in real time, and generate a complete report with treatment recommendations." if st.session_state.language == 'en' else "Esta plataforma utiliza redes neurais convolucionais (CNN) avançadas para identificar com precisão e rapidez patologias comuns em folhas de videira. Carregue uma imagem, compare as previsões de múltiplos modelos em tempo real e gere um relatório completo com recomendações de tratamento."}
            </p>
        </div>
        """, unsafe_allow_html=True)
        
        # Tarjetas de características
        features_html = ""
        if st.session_state.language == 'es':
            features_html = """
            <div class="features-grid">
                <div class="feature-card">
                    <span class="feature-icon">🔍</span>
                    <div class="feature-content">
                        <h4>1. Diagnóstico por imagen</h4>
                        <p>Analice instantáneamente fotos de hojas de vid subidas o tomadas con la cámara en campo.</p>
                    </div>
                </div>
                <div class="feature-card">
                    <span class="feature-icon">🧠</span>
                    <div class="feature-content">
                        <h4>2. Modelos Clásicos e Híbridos</h4>
                        <p>5 modelos: SVM, Random Forest, KNN (clásicos) y CNN+SVM, MobileNetV2+RF (híbridos).</p>
                    </div>
                </div>
                <div class="feature-card">
                    <span class="feature-icon">📊</span>
                    <div class="feature-content">
                        <h4>3. Validación estadística</h4>
                        <p>Evaluación matemática mediante métricas MCC y contrastes con pruebas de McNemar.</p>
                    </div>
                </div>
                <div class="feature-card">
                    <span class="feature-icon">📄</span>
                    <div class="feature-content">
                        <h4>4. Reporte PDF automático</h4>
                        <p>Descargue un reporte técnico estructurado con diagnósticos y medidas de prevención fitosanitaria.</p>
                    </div>
                </div>
            </div>
            """
        elif st.session_state.language == 'en':
            features_html = """
            <div class="features-grid">
                <div class="feature-card">
                    <span class="feature-icon">🔍</span>
                    <div class="feature-content">
                        <h4>1. Image Diagnosis</h4>
                        <p>Instantly analyze grapevine leaf photos uploaded or captured directly in the field.</p>
                    </div>
                </div>
                <div class="feature-card">
                    <span class="feature-icon">🧠</span>
                    <div class="feature-content">
                        <h4>2. Classic &amp; Hybrid Models</h4>
                        <p>5 models: SVM, Random Forest, KNN (classic) and CNN+SVM, MobileNetV2+RF (hybrid).</p>
                    </div>
                </div>
                <div class="feature-card">
                    <span class="feature-icon">📊</span>
                    <div class="feature-content">
                        <h4>3. Statistical Validation</h4>
                        <p>Mathematical evaluation utilizing MCC metrics and McNemar pairwise hypothesis tests.</p>
                    </div>
                </div>
                <div class="feature-card">
                    <span class="feature-icon">📄</span>
                    <div class="feature-content">
                        <h4>4. Automatic PDF Report</h4>
                        <p>Download a structured technical report containing diagnosis and phytosanitary action plans.</p>
                    </div>
                </div>
            </div>
            """
        else: # pt
            features_html = """
            <div class="features-grid">
                <div class="feature-card">
                    <span class="feature-icon">🔍</span>
                    <div class="feature-content">
                        <h4>1. Diagnóstico por Imagem</h4>
                        <p>Analise instantaneamente fotos de folhas de videira carregadas ou tiradas no campo.</p>
                    </div>
                </div>
                <div class="feature-card">
                    <span class="feature-icon">🧠</span>
                    <div class="feature-content">
                        <h4>2. Modelos Clássicos e Híbridos</h4>
                        <p>5 modelos: SVM, Random Forest, KNN (clássicos) e CNN+SVM, MobileNetV2+RF (híbridos).</p>
                    </div>
                </div>
                <div class="feature-card">
                    <span class="feature-icon">📊</span>
                    <div class="feature-content">
                        <h4>3. Validação Estatística</h4>
                        <p>Avaliação matemática usando métricas MCC e contrastes com testes de McNemar.</p>
                    </div>
                </div>
                <div class="feature-card">
                    <span class="feature-icon">📄</span>
                    <div class="feature-content">
                        <h4>4. Relatório PDF Automático</h4>
                        <p>Baixe um relatório técnico estruturado com diagnósticos e ações de prevenção fitossanitária.</p>
                    </div>
                </div>
            </div>
            """
            
        st.markdown(features_html, unsafe_allow_html=True)
        
        # Mensaje de advertencia / instrucciones para cargar modelos
        warning_title = "Para iniciar el diagnóstico, cargue los modelos entrenados desde la barra lateral." if st.session_state.language == 'es' else "To start the diagnosis, please load the trained models from the sidebar." if st.session_state.language == 'en' else "Para iniciar o diagnóstico, carregue os modelos treinados na barra lateral."
        warning_text = "Utilice el botón principal 'Cargar Modelos' para inicializar las redes neuronales en memoria." if st.session_state.language == 'es' else "Use the main 'Load Models' button to initialize the neural networks in memory." if st.session_state.language == 'en' else "Use o botão principal 'Carregar Modelos' para inicializar as redes neurais em memória."
        warning_note = "Una vez cargados, podrá subir imágenes de hojas de vid para obtener el diagnóstico y recomendaciones." if st.session_state.language == 'es' else "Once loaded, you will be able to upload vine leaf images to obtain diagnoses and recommendations." if st.session_state.language == 'en' else "Uma vez carregados, você poderá enviar imagens de folhas de videira para obter o diagnóstico e as recomendações."
        
        st.markdown(f"""
        <div class="load-warning-card">
            <h4 class="load-warning-title">⚠️ {warning_title}</h4>
            <p class="load-warning-text">{warning_text}</p>
            <p class="load-warning-note">{warning_note}</p>
        </div>
        """, unsafe_allow_html=True)
        
        return

    # Tabs principales
    tab1, tab2, tab3, tab4 = st.tabs([
        get_text('tab_diagnosis', st.session_state.language), 
        get_text('tab_statistical', st.session_state.language), 
        get_text('tab_validation', st.session_state.language), 
        get_text('tab_info', st.session_state.language)
    ])

    with tab1:
        st.header(get_text('diagnosis_title', st.session_state.language))

        # Opciones de entrada
        col1, col2 = st.columns([2, 1])
        with col1:
            input_method = st.radio(
                get_text('input_method', st.session_state.language),
                [get_text('upload_image', st.session_state.language), get_text('use_camera', st.session_state.language)],
                horizontal=True
            )

        # Subir imagen
        if input_method == get_text('upload_image', st.session_state.language):
            uploaded_file = st.file_uploader(
                get_text('select_image', st.session_state.language),
                type=['jpg', 'jpeg', 'png'],
                help=get_text('supported_formats', st.session_state.language)
            )

            if uploaded_file is not None:
                # Cargar y mostrar imagen
                image = Image.open(uploaded_file).convert('RGB')
                st.session_state.current_image = image

                # Mostrar imagen
                col1, col2 = st.columns([1, 1])
                with col1:
                    st.image(image, caption=get_text('image_loaded', st.session_state.language), use_column_width=True)

                # Botón de análisis
                if st.button(get_text('analyze_image', st.session_state.language), type="primary"):
                    with st.spinner(get_text('analyzing', st.session_state.language)):
                        # Realizar predicciones con todos los modelos
                        results = []
                        for model_name, model in st.session_state.models.items():
                            result = predict_disease(image, model, model_name)
                            results.append(result)

                        st.session_state.predictions = results
                        st.session_state.pdf_bytes = None
                        st.session_state.pdf_ready = False

                # Mostrar resultados si existen
                if st.session_state.predictions:
                    st.success(get_text('analysis_completed', st.session_state.language))

                    # Mostrar resultados por modelo
                    st.subheader(get_text('diagnosis_results', st.session_state.language))

                    # Crear columnas para cada modelo
                    cols = st.columns(len(st.session_state.predictions))

                    for i, result in enumerate(st.session_state.predictions):
                        with cols[i]:
                            # Métrica principal
                            st.metric(
                                label=result['model_name'],
                                value=result['predicted_class_es'],
                                delta=f"{result['confidence']:.1%} {get_text('confidence', st.session_state.language)}"
                            )
                            st.caption(f"⏱️ {result['inference_time']:.1f} ms")

                    # Consenso de modelos
                    st.subheader(get_text('consensus_diagnosis', st.session_state.language))

                    # Calcular diagnóstico más frecuente
                    predictions = [r['predicted_class'] for r in st.session_state.predictions]
                    consensus = max(set(predictions), key=predictions.count)
                    consensus_count = predictions.count(consensus)

                    # Calcular confianza promedio para el consenso
                    consensus_confidence = np.mean([
                        r['confidence'] for r in st.session_state.predictions
                        if r['predicted_class'] == consensus
                    ])

                    # Mostrar consenso
                    col1, col2, col3 = st.columns([2, 1, 1])
                    with col1:
                        st.info(f"**{get_text('final_diagnosis', st.session_state.language)}** {get_disease_names(st.session_state.language)[consensus]}")
                    with col2:
                        st.metric(get_text('coincidence', st.session_state.language), f"{consensus_count}/{len(predictions)}")
                    with col3:
                        st.metric(get_text('confidence', st.session_state.language).title(), f"{consensus_confidence:.1%}")

                    # Gráfico de probabilidades
                    st.subheader(get_text('probability_distribution', st.session_state.language))

                    # Preparar datos para el gráfico
                    fig, axes = plt.subplots(1, len(st.session_state.predictions),
                                             figsize=(12, 4))
                    if len(st.session_state.predictions) == 1:
                        axes = [axes]

                    for i, (ax, result) in enumerate(zip(axes, st.session_state.predictions)):
                        probs = result['all_predictions']
                        ax.barh(DISEASE_CLASSES, probs, color=['#e74c3c', '#f39c12', '#27ae60', '#3498db'])
                        ax.set_xlim(0, 1)
                        ax.set_title(result['model_name'])
                        ax.set_xlabel('Probabilidad')

                        # Añadir valores en las barras
                        for j, (clase, prob) in enumerate(zip(DISEASE_CLASSES, probs)):
                            ax.text(prob + 0.02, j, f'{prob:.1%}',
                                    va='center', fontsize=9)

                    plt.tight_layout()
                    st.pyplot(fig)

                    # Recomendaciones
                    st.subheader(get_text('treatment_recommendations', st.session_state.language))
                    recommendations = get_treatment_recommendations(consensus, st.session_state.language)

                    if recommendations:
                        # Título y gravedad
                        col1, col2 = st.columns([3, 1])
                        with col1:
                            st.markdown(f"### {recommendations['titulo']}")
                        with col2:
                            if recommendations['gravedad'] == "Alta":
                                st.error(f"Gravedad: {recommendations['gravedad']}")
                            elif recommendations['gravedad'] == "Muy Alta":
                                st.error(f"Gravedad: {recommendations['gravedad']}")
                            elif recommendations['gravedad'] == "Moderada":
                                st.warning(f"Gravedad: {recommendations['gravedad']}")
                            else:
                                st.success(f"Gravedad: {recommendations['gravedad']}")

                        # Tratamiento
                        with st.expander(get_text('recommended_treatment', st.session_state.language), expanded=True):
                            for item in recommendations['tratamiento']:
                                st.write(f"• {item}")

                        # Prevención
                        with st.expander(get_text('preventive_measures', st.session_state.language)):
                            for item in recommendations['prevencion']:
                                st.write(f"• {item}")

                    # Mostrar ranking de modelos
                    if st.session_state.ranking_data:
                        st.markdown("---")
                        st.subheader(get_text('model_ranking', st.session_state.language))
                        ranking_df = pd.DataFrame(st.session_state.ranking_data)
                        if not ranking_df.empty and 'modelo' in ranking_df.columns:
                            display_cols = [c for c in ['ranking', 'modelo', 'accuracy', 'f1_score', 'mcc'] if c in ranking_df.columns]
                            if display_cols:
                                st.dataframe(ranking_df[display_cols].head(10), use_container_width=True)

                    # Matrices de confusión y curvas ROC
                    cm_dir = Path("reports/modelos")
                    cm_files = sorted(cm_dir.glob("confusion_*.png"))
                    roc_files = sorted(cm_dir.glob("roc_*.png"))
                    if cm_files or roc_files:
                        st.markdown("---")
                        st.subheader("📊 Matrices de Confusión")
                        cm_cols = st.columns(min(len(cm_files), 5))
                        for i, f in enumerate(cm_files):
                            with cm_cols[i % len(cm_cols)]:
                                model_label = f.stem.replace("confusion_", "").replace("_", " ").title()
                                st.caption(model_label)
                                st.image(str(f), use_column_width=True)
                    if roc_files:
                        st.subheader("📈 Curvas ROC (One-vs-Rest)")
                        roc_cols = st.columns(min(len(roc_files), 5))
                        for i, f in enumerate(roc_files):
                            with roc_cols[i % len(roc_cols)]:
                                model_label = f.stem.replace("roc_", "").replace("_", " ").title()
                                st.caption(model_label)
                                st.image(str(f), use_column_width=True)

                    # Botones para generar reportes
                    st.subheader(get_text('generate_report', st.session_state.language))
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        if st.button(get_text('download_pdf', st.session_state.language)):
                            with st.spinner(get_text('generating_report', st.session_state.language)):
                                st.session_state.pdf_bytes = generate_diagnosis_pdf(
                                    image,
                                    st.session_state.predictions,
                                    consensus
                                )
                                st.session_state.pdf_ready = True
                            st.rerun()
                    with col2:
                        if st.button(get_text('download_word', st.session_state.language)):
                            with st.spinner(get_text('generating_report', st.session_state.language)):
                                st.session_state.docx_bytes = generate_diagnosis_docx(
                                    image,
                                    st.session_state.predictions,
                                    consensus
                                )
                                st.session_state.docx_ready = True
                            st.rerun()
                    with col3:
                        if st.button(get_text('download_excel', st.session_state.language)):
                            with st.spinner(get_text('generating_report', st.session_state.language)):
                                st.session_state.xlsx_bytes = generate_diagnosis_xlsx(
                                    image,
                                    st.session_state.predictions,
                                    consensus
                                )
                                st.session_state.xlsx_ready = True
                            st.rerun()

                    # Botones de descarga
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        if st.session_state.get("pdf_ready") and st.session_state.get("pdf_bytes"):
                            st.download_button(
                                label=get_text('download_pdf_button', st.session_state.language),
                                data=st.session_state.pdf_bytes,
                                file_name=f"diagnostico_vineguard_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                                mime="application/pdf"
                            )
                    with col2:
                        if st.session_state.get("docx_ready") and st.session_state.get("docx_bytes"):
                            st.download_button(
                                label=get_text('download_word_button', st.session_state.language),
                                data=st.session_state.docx_bytes,
                                file_name=f"diagnostico_vineguard_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                    with col3:
                        if st.session_state.get("xlsx_ready") and st.session_state.get("xlsx_bytes"):
                            st.download_button(
                                label=get_text('download_excel_button', st.session_state.language),
                                data=st.session_state.xlsx_bytes,
                                file_name=f"diagnostico_vineguard_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                            )

        else:  # Usar cámara
            st.info(get_text('camera_info', st.session_state.language))
            st.warning(get_text('camera_warning', st.session_state.language))

    with tab2:
        st.header(get_text('tab_statistical', st.session_state.language))

        # Verificar si hay análisis de validación real disponible
        if st.session_state.mcnemar_analysis and st.session_state.mcnemar_analysis.get('real_data', False):
            # Mostrar análisis real de múltiples imágenes
            analysis = st.session_state.mcnemar_analysis

            st.success("✅ **Análisis con datos reales disponible** (de validación McNemar)")

            # Coeficiente de Matthews REAL
            st.subheader("📈 Coeficiente de Matthews (MCC) - Datos Reales")

            st.markdown("""
            <div class="statistical-box" style="color: black;">
            <h4 style="color: black;">🧮 ¿Qué es el Coeficiente de Matthews?</h4>
            <p>El MCC es una métrica balanceada que considera todos los tipos de predicciones (verdaderos/falsos positivos/negativos). 
            Valores cercanos a +1 indican predicción perfecta, 0 indica predicción aleatoria, y -1 indica predicción completamente incorrecta.</p>
            </div>
            """, unsafe_allow_html=True)

            # Mostrar MCC para cada modelo
            col1, col2 = st.columns([2, 1])

            with col1:
                # Tabla de MCC
                mcc_data = []
                for mcc_result in analysis['matthews_coefficients']:
                    mcc_data.append({
                        'Modelo': mcc_result['model'],
                        'MCC': f"{mcc_result['mcc']:.3f}",
                        'Interpretación': mcc_result['interpretation']
                    })

                mcc_df = pd.DataFrame(mcc_data)
                st.table(mcc_df)

            with col2:
                # Gráfico de MCC
                fig, ax = plt.subplots(figsize=(6, 4))
                models = [m['model'] for m in analysis['matthews_coefficients']]
                mccs = [m['mcc'] for m in analysis['matthews_coefficients']]

                bars = ax.bar(models, mccs, color=['#3498db', '#2ecc71', '#e74c3c', '#f39c12'])
                ax.set_ylabel('Coeficiente de Matthews')
                ax.set_title('MCC por Modelo (Datos Reales)')
                ax.set_ylim(-1, 1)
                ax.axhline(y=0, color='black', linestyle='--', alpha=0.3)

                # Añadir valores en las barras
                for bar, mcc in zip(bars, mccs):
                    height = bar.get_height()
                    ax.text(bar.get_x() + bar.get_width()/2., height + 0.02,
                            f'{mcc:.3f}', ha='center', va='bottom')

                plt.xticks(rotation=45)
                plt.tight_layout()
                st.pyplot(fig)

            # Comparación general
            st.subheader("🏆 Ranking de Modelos")

            # Ordenar modelos por MCC
            mcc_sorted = sorted(analysis['matthews_coefficients'], key=lambda x: x['mcc'], reverse=True)

            st.write("**Ranking basado en Coeficiente de Matthews (Datos Reales):**")
            for i, model_result in enumerate(mcc_sorted):
                if i == 0:
                    st.success(f"🥇 **1º lugar:** {model_result['model']} (MCC: {model_result['mcc']:.3f})")
                elif i == 1:
                    st.info(f"🥈 **2º lugar:** {model_result['model']} (MCC: {model_result['mcc']:.3f})")
                elif i == 2:
                    st.warning(f"🥉 **3º lugar:** {model_result['model']} (MCC: {model_result['mcc']:.3f})")
                else:
                    st.write(f"**{i+1}º lugar:** {model_result['model']} (MCC: {model_result['mcc']:.3f})")

            # Información del dataset usado
            st.info(f"**Tamaño de muestra:** {analysis['sample_size']} imágenes reales")

        # Si tenemos predicciones de una imagen, mostrar solo análisis de velocidad
        elif st.session_state.predictions:
            st.subheader("⚡ Análisis de Velocidad de Modelos")

            # Obtener datos de velocidad
            model_names = [result['model_name'] for result in st.session_state.predictions]
            inference_times = [result['inference_time'] for result in st.session_state.predictions]

            # Crear gráfico circular
            fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 5))

            # Gráfico circular de distribución de tiempos
            colors = ['#3498db', '#2ecc71', '#e74c3c', '#f39c12'][:len(model_names)]
            wedges, texts, autotexts = ax1.pie(inference_times,
                                               labels=model_names,
                                               autopct='%1.1f ms',
                                               colors=colors,
                                               startangle=90)
            ax1.set_title('Distribución de Tiempos de Inferencia')

            # Hacer el texto más legible
            for autotext in autotexts:
                autotext.set_color('white')
                autotext.set_fontweight('bold')

            # Gráfico de barras comparativo
            bars = ax2.bar(range(len(model_names)), inference_times, color=colors)
            ax2.set_xlabel('Modelos')
            ax2.set_ylabel('Tiempo (ms)')
            ax2.set_title('Comparación de Velocidad')
            ax2.set_xticks(range(len(model_names)))
            ax2.set_xticklabels([name.replace(' ', '\n') for name in model_names], rotation=0)

            # Añadir valores en las barras
            for bar, time in zip(bars, inference_times):
                height = bar.get_height()
                ax2.text(bar.get_x() + bar.get_width()/2., height + 0.5,
                         f'{time:.1f}ms', ha='center', va='bottom', fontweight='bold')

            plt.tight_layout()
            st.pyplot(fig)

            # Métricas de velocidad
            col1, col2, col3 = st.columns(3)

            with col1:
                fastest_idx = np.argmin(inference_times)
                st.success(f"**🚀 Más Rápido**\n{model_names[fastest_idx]}\n{inference_times[fastest_idx]:.1f} ms")

            with col2:
                slowest_idx = np.argmax(inference_times)
                st.error(f"**🐌 Más Lento**\n{model_names[slowest_idx]}\n{inference_times[slowest_idx]:.1f} ms")

            with col3:
                avg_time = np.mean(inference_times)
                st.info(f"**⏱️ Promedio**\nTodos los modelos\n{avg_time:.1f} ms")

            # Estadísticas adicionales de velocidad
            st.markdown("**📈 Estadísticas de Velocidad:**")
            speed_stats = pd.DataFrame({
                'Modelo': model_names,
                'Tiempo (ms)': [f"{t:.1f}" for t in inference_times],
                'Velocidad Relativa': [f"{(min(inference_times)/t)*100:.1f}%" for t in inference_times],
                'Diferencia vs Más Rápido': [f"+{t-min(inference_times):.1f} ms" if t != min(inference_times) else "Baseline" for t in inference_times]
            })
            st.table(speed_stats)

            # Nota sobre análisis estadístico
            st.warning("""
            ⚠️ **Análisis Estadístico No Disponible**
            
            Para obtener análisis estadístico real (MCC y McNemar):
            1. Ve a la pestaña '🔬 Validación McNemar'
            2. Carga al menos 30 imágenes con sus etiquetas verdaderas
            3. El análisis estadístico aparecerá automáticamente aquí
            
            **¿Por qué necesitas múltiples imágenes?**
            - Con una sola imagen no se pueden calcular métricas estadísticas reales
            - Se requieren al menos 30 muestras para resultados confiables
            - MCC y McNemar comparan el rendimiento general de los modelos
            """)

        else:
            # No hay datos disponibles
            st.info("👆 Realiza un diagnóstico o validación para generar el análisis estadístico")

            # Mostrar información sobre las pruebas estadísticas
            st.subheader("📚 Acerca de las Pruebas Estadísticas")

            col1, col2 = st.columns(2)

            with col1:
                st.markdown("""
                **🧮 Coeficiente de Matthews (MCC)**
                
                - Métrica balanceada para clasificación
                - Rango: -1 (peor) a +1 (mejor)
                - Considera todos los tipos de predicción
                - Útil para datasets desbalanceados
                - Interpretación:
                  - MCC ≥ 0.8: Muy bueno
                  - MCC ≥ 0.6: Bueno  
                  - MCC ≥ 0.4: Moderado
                  - MCC < 0.4: Necesita mejora
                """)

            with col2:
                st.markdown("""
                **🔬 Prueba de McNemar**
                
                - Compara dos modelos estadísticamente
                - Basada en distribución χ² (chi-cuadrado)
                - H₀: No hay diferencia entre modelos
                - H₁: Hay diferencia significativa
                - Interpretación del p-valor:
                  - p < 0.001: Muy significativo
                  - p < 0.01: Significativo
                  - p < 0.05: Marginalmente significativo
                  - p ≥ 0.05: No significativo
                """)

    with tab3:
        st.header("🔬 Validación Estadística con Dataset Real")

        if not st.session_state.models_loaded:
            st.warning("👈 Por favor, carga los modelos desde la barra lateral primero")
        else:
            # ====== TEORÍA AL INICIO ======
            st.markdown("### 📚 Fundamentos Teóricos")

            col1, col2 = st.columns(2)

            with col1:
                st.markdown("""
                <div class="theory-box">
                <h4>🧮 Coeficiente de Matthews (MCC)</h4>
                <p><strong>Fórmula:</strong> MCC = (TP×TN - FP×FN) / √[(TP+FP)(TP+FN)(TN+FP)(TN+FN)]</p>
                <p><strong>Propósito:</strong> Métrica balanceada que evalúa la calidad general de clasificación considerando todas las categorías de predicción.</p>
                <p><strong>Ventajas:</strong> Robusto ante clases desbalanceadas, interpretación intuitiva (-1 a +1), y considera todos los aspectos de la matriz de confusión.</p>
                </div>
                """, unsafe_allow_html=True)

            with col2:
                st.markdown("""
                <div class="theory-box">
                <h4>🔬 Prueba de McNemar</h4>
                <p><strong>Fórmula:</strong> χ² = (|b - c| - 0.5)² / (b + c)</p>
                <p><strong>Propósito:</strong> Test estadístico que compara el rendimiento de dos clasificadores para determinar si sus diferencias son significativas.</p>
                <p><strong>Aplicación:</strong> Validación científica de que un modelo es estadísticamente superior a otro (p < 0.05 = diferencia significativa).</p>
                </div>
                """, unsafe_allow_html=True)

            st.markdown("---")

            # ====== INTERFAZ DINÁMICA CON CARPETAS ======
            st.markdown("""
            **📁 Sistema de Validación por Carpetas Inteligentes**
            
            📋 **Instrucciones:**
            - Organiza tus imágenes por enfermedad en cada "carpeta" digital
            - Mínimo recomendado: 30+ imágenes totales (10+ por categoría)
            - El sistema automáticamente etiquetará las imágenes según la carpeta elegida
            """)

            st.subheader("🗂️ Carpetas de Enfermedades")

            # Crear las 4 carpetas dinámicas
            disease_files = {}

            # Layout en grid 2x2
            row1_col1, row1_col2 = st.columns(2)
            row2_col1, row2_col2 = st.columns(2)

            columns = [row1_col1, row1_col2, row2_col1, row2_col2]
            disease_folders = get_disease_folders(st.session_state.language)
            disease_names = list(disease_folders.keys())

            for i, (disease_name, col) in enumerate(zip(disease_names, columns)):
                with col:
                    folder_info = disease_folders[disease_name]

                    st.markdown(f"""
                    <div class="disease-folder {folder_info['css_class']}">
                    <h4 style="text-align: center; margin-bottom: 10px;">
                    {folder_info['icon']} {disease_name}
                    </h4>
                    <p style="text-align: center; font-size: 0.9em; margin-bottom: 15px;">
                    {folder_info['description']}
                    </p>
                    </div>
                    """, unsafe_allow_html=True)

                    # File uploader para cada enfermedad
                    uploaded_files = st.file_uploader(
                        f"Subir imágenes de {disease_name}",
                        type=['jpg', 'jpeg', 'png'],
                        accept_multiple_files=True,
                        key=f"files_{disease_name}",
                        help=f"Arrastra aquí las imágenes de {disease_name}"
                    )

                    if uploaded_files:
                        disease_files[disease_name] = uploaded_files
                        st.success(f"✅ {len(uploaded_files)} imágenes cargadas")
                    else:
                        disease_files[disease_name] = []

            # ====== RESUMEN DEL DATASET ======
            total_images = sum(len(files) for files in disease_files.values())

            if total_images > 0:
                st.markdown("---")
                st.subheader("📊 Resumen del Dataset")

                # Mostrar distribución
                col1, col2 = st.columns([1, 1])

                with col1:
                    st.markdown("**Distribución por enfermedad:**")
                    for disease_name, files in disease_files.items():
                        if len(files) > 0:
                            icon = disease_folders[disease_name]["icon"]
                            st.write(f"{icon} **{disease_name}:** {len(files)} imágenes")

                    st.markdown(f"**📈 Total:** {total_images} imágenes")

                    # Recomendaciones
                    if total_images < 30:
                        st.warning("⚠️ Se recomienda al menos 30 imágenes para resultados estadísticamente válidos")
                    else:
                        st.success("✅ Dataset suficiente para análisis estadístico robusto")

                with col2:
                    # Gráfico de distribución
                    if total_images > 0:
                        labels = []
                        sizes = []
                        colors = []

                        color_map = {
                            "Black_rot": "#e74c3c",
                            "Esca": "#8B4513",
                            "Healthy": "#27ae60",
                            "Leaf_blight": "#f39c12"
                        }

                        for disease_name, files in disease_files.items():
                            if len(files) > 0:
                                labels.append(disease_name.replace(" ", "\n"))
                                sizes.append(len(files))
                                # Obtener la clave interna (inglés) para mapear el color de forma consistente en cualquier idioma
                                internal_key = disease_folders[disease_name]["key"]
                                colors.append(color_map.get(internal_key, "#cccccc"))

                        if sizes:
                            fig, ax = plt.subplots(figsize=(6, 4))
                            wedges, texts, autotexts = ax.pie(sizes, labels=labels, autopct='%1.0f%%',
                                                              colors=colors, startangle=90)
                            ax.set_title('Distribución del Dataset', fontweight='bold')

                            # Mejorar legibilidad
                            for autotext in autotexts:
                                autotext.set_color('white')
                                autotext.set_fontweight('bold')

                            plt.tight_layout()
                            st.pyplot(fig)

                # ====== BOTÓN DE PROCESAMIENTO ======
                st.markdown("---")

                col1, col2, col3 = st.columns([0.2, 4.6, 0.2])

                with col2:
                    if st.button("🚀 PROCESAR DATASET Y CALCULAR ESTADÍSTICAS", type="primary", use_container_width=True):
                        with st.spinner("🔄 Procesando imágenes y realizando análisis estadístico..."):

                            # Procesar imágenes por carpetas
                            validation_data, error = process_multiple_images_by_folders(
                                disease_files, st.session_state.models
                            )

                            if error:
                                st.error(f"❌ Error: {error}")
                            else:
                                # Calcular estadísticas con datos reales
                                mcnemar_analysis = perform_mcnemar_analysis(validation_data)

                                # Guardar en session_state para uso posterior
                                st.session_state.mcnemar_validation = validation_data
                                st.session_state.mcnemar_analysis = mcnemar_analysis

                                # ====== MOSTRAR RESULTADOS DESTACADOS ======
                                st.markdown("""
                                <div class="result-highlight">
                                <h2 style="color: white; text-align: center; margin-bottom: 20px;">
                                ✅ ¡ANÁLISIS ESTADÍSTICO COMPLETADO!
                                </h2>
                                <p style="color: white; text-align: center; font-size: 1.2em;">
                                Datos procesados con éxito. Resultados científicamente válidos generados.
                                </p>
                                </div>
                                """, unsafe_allow_html=True)

                                # ====== RESULTADOS DE VALIDACIÓN ======
                                st.subheader("📊 Resultados de Validación")

                                # Tabla de precisión por modelo
                                results_df = create_validation_results_display(validation_data, mcnemar_analysis)
                                st.write("**Precisión por modelo:**")

                                # Colorear la tabla
                                styled_df = results_df.style.apply(lambda x: ['background-color: #000000' if i == 0 else '' for i in range(len(x))], axis=0)
                                st.dataframe(styled_df, use_container_width=True)

                                # ====== MCC CON VISUALIZACIÓN MEJORADA ======
                                st.subheader("📈 Coeficiente de Matthews (MCC) - Análisis Real")

                                col1, col2 = st.columns([2, 1])
                                with col1:
                                    mcc_data = []
                                    for mcc_result in mcnemar_analysis['matthews_coefficients']:
                                        mcc_data.append({
                                            'Modelo': mcc_result['model'],
                                            'MCC': f"{mcc_result['mcc']:.3f}",
                                            'Interpretación': mcc_result['interpretation']
                                        })
                                    mcc_df = pd.DataFrame(mcc_data)
                                    st.table(mcc_df)

                                with col2:
                                    # Gráfico de MCC mejorado
                                    fig, ax = plt.subplots(figsize=(6, 4))
                                    models = [m['model'] for m in mcnemar_analysis['matthews_coefficients']]
                                    mccs = [m['mcc'] for m in mcnemar_analysis['matthews_coefficients']]

                                    bars = ax.bar(models, mccs, color=['#3498db', '#2ecc71', '#e74c3c', '#f39c12'])
                                    ax.set_ylabel('Coeficiente de Matthews', fontweight='bold')
                                    ax.set_title('MCC por Modelo', fontweight='bold', fontsize=14)
                                    ax.set_ylim(-1, 1)
                                    ax.axhline(y=0, color='black', linestyle='--', alpha=0.3)
                                    ax.grid(True, alpha=0.3)

                                    for bar, mcc in zip(bars, mccs):
                                        height = bar.get_height()
                                        ax.text(bar.get_x() + bar.get_width()/2., height + 0.02,
                                                f'{mcc:.3f}', ha='center', va='bottom', fontweight='bold')

                                    plt.xticks(rotation=45)
                                    plt.tight_layout()
                                    st.pyplot(fig)

                                # ====== RESULTADOS DE MCNEMAR COMPACTOS ======
                                st.subheader("🔬 Resultados de la Prueba de McNemar")

                                # Resumen ejecutivo de McNemar
                                significant_count = len([r for r in mcnemar_analysis['mcnemar_results'] if r['p_value'] < 0.05])

                                if significant_count > 0:
                                    st.warning(f"⚠️ **{significant_count} de {len(mcnemar_analysis['mcnemar_results'])} comparaciones muestran diferencias significativas**")
                                else:
                                    st.success(f"✅ **Ninguna diferencia significativa encontrada** entre los {len(mcnemar_analysis['mcnemar_results'])} pares de modelos")

                                # Mostrar comparaciones en formato compacto
                                for mcnemar_result in mcnemar_analysis['mcnemar_results']:
                                    col1, col2, col3, col4 = st.columns([2, 1, 1, 2])

                                    with col1:
                                        st.write(f"**{mcnemar_result['model1']}** vs **{mcnemar_result['model2']}**")
                                    with col2:
                                        st.metric("χ²", f"{mcnemar_result['statistic']:.3f}")
                                    with col3:
                                        st.metric("p-valor", f"{mcnemar_result['p_value']:.4f}")
                                    with col4:
                                        if mcnemar_result['p_value'] < 0.05:
                                            st.error("**Significativo**")
                                        else:
                                            st.success("**No significativo**")

                                # ====== INTERPRETACIÓN PARA EL PROFESOR ======
                                interpretation = generate_interpretation_for_professor(mcnemar_analysis, validation_data)

                                st.markdown("""
                                <div class="interpretation-box">
                                {}
                                </div>
                                """.format(interpretation.replace('\n', '<br>')), unsafe_allow_html=True)

                                # ====== ENLACE A ANÁLISIS COMPLETO ======
                                st.info("""
                                ✅ **Los resultados completos están disponibles en la pestaña 'Análisis Estadístico'**
                                
                                Ve a la pestaña anterior para explorar visualizaciones detalladas y métricas adicionales.
                                """)

            else:
                st.info("📁 Carga imágenes en las carpetas de enfermedades para comenzar el análisis estadístico")

    with tab4:
        st.header("📚 Información sobre Enfermedades")

        # Información detallada de cada enfermedad
        disease_info = {
            "Podredumbre Negra (Black Rot)": {
                "descripcion": "Causada por el hongo Guignardia bidwellii. Una de las enfermedades más destructivas de la vid.",
                "sintomas": [
                    "Manchas circulares marrones en las hojas",
                    "Lesiones negras en los frutos",
                    "Momificación de las bayas",
                    "Picnidios negros en tejidos infectados"
                ],
                "condiciones": "Se desarrolla en condiciones de alta humedad y temperaturas de 20-27°C",
                "imagen": "🔴"
            },
            "Esca (Sarampión Negro)": {
                "descripcion": "Enfermedad compleja causada por varios hongos. Afecta el sistema vascular de la planta.",
                "sintomas": [
                    "Decoloración intervenal en las hojas",
                    "Necrosis marginal",
                    "Muerte regresiva de brotes",
                    "Pudrición interna del tronco"
                ],
                "condiciones": "Se agrava con estrés hídrico y heridas de poda mal protegidas",
                "imagen": "🟤"
            },
            "Tizón de la Hoja (Leaf Blight)": {
                "descripcion": "Causada por el hongo Isariopsis. Afecta principalmente las hojas maduras.",
                "sintomas": [
                    "Manchas angulares amarillentas",
                    "Necrosis foliar progresiva",
                    "Defoliación prematura",
                    "Reducción del vigor de la planta"
                ],
                "condiciones": "Favorecida por alta humedad relativa y temperaturas moderadas",
                "imagen": "🟡"
            }
        }

        for disease_name, info in disease_info.items():
            with st.expander(f"{info['imagen']} {disease_name}"):
                st.write(f"**Descripción:** {info['descripcion']}")

                st.write("**Síntomas:**")
                for sintoma in info['sintomas']:
                    st.write(f"• {sintoma}")

                st.write(f"**Condiciones favorables:** {info['condiciones']}")

        # Buenas prácticas
        st.subheader("✅ Buenas Prácticas de Manejo")

        col1, col2 = st.columns(2)

        with col1:
            st.markdown("""
            **Prevención:**
            - Monitoreo regular del viñedo
            - Poda sanitaria adecuada
            - Manejo del dosel vegetal
            - Drenaje apropiado del suelo
            - Selección de variedades resistentes
            """)

        with col2:
            st.markdown("""
            **Manejo Integrado:**
            - Uso racional de fungicidas
            - Rotación de ingredientes activos
            - Aplicaciones en momentos críticos
            - Registro de aplicaciones
            - Evaluación de eficacia
            """)

        # Sección de ranking y mejor modelo
        if st.session_state.ranking_data:
            st.markdown("---")
            st.subheader("🏆 Ranking de Modelos")
            ranking_df = pd.DataFrame(st.session_state.ranking_data)
            if not ranking_df.empty and 'modelo' in ranking_df.columns:
                display_cols = [c for c in ['ranking', 'modelo', 'accuracy', 'f1_score', 'mcc'] if c in ranking_df.columns]
                if display_cols:
                    st.dataframe(ranking_df[display_cols], use_container_width=True)
                mejor = ranking_df.iloc[0] if 'ranking' in ranking_df.columns else ranking_df.iloc[0]
                st.success(f"**Modelo recomendado:** {mejor['modelo']}")

        # Información sobre pruebas estadísticas
        st.subheader("📊 Sobre las Pruebas Estadísticas")

        with st.expander("🧮 Coeficiente de Matthews - Información Técnica"):
            st.markdown("""
            **Fórmula del MCC:**
            
            MCC = (TP×TN - FP×FN) / √[(TP+FP)(TP+FN)(TN+FP)(TN+FN)]
            
            Donde:
            - TP = Verdaderos Positivos
            - TN = Verdaderos Negativos  
            - FP = Falsos Positivos
            - FN = Falsos Negativos
            
            **Ventajas:**
            - Balanceado para todas las clases
            - Robusto ante datasets desbalanceados
            - Fácil interpretación (-1 a +1)
            - Considera todos los aspectos de la matriz de confusión
            """)

        with st.expander("🔬 Prueba de McNemar - Información Técnica"):
            st.markdown("""
            **Procedimiento:**
            
            1. **Hipótesis:**
               - H₀: No hay diferencia entre modelos
               - H₁: Hay diferencia significativa
            
            2. **Estadístico de prueba:**
               χ² = (|b - c| - 0.5)² / (b + c)
               
               Donde b y c son las frecuencias de desacuerdo entre modelos
            
            3. **Decisión:**
               - Si p < 0.05: Rechazar H₀ (hay diferencia)
               - Si p ≥ 0.05: No rechazar H₀ (sin diferencia)
            
            **Aplicación:**
            - Comparación objetiva de modelos
            - Base estadística para selección de modelos
            - Validación de mejoras en algoritmos
            """)

        # Calendario de aplicaciones
        st.subheader("📅 Calendario de Protección Fitosanitaria")

        calendar_data = {
            "Etapa Fenológica": ["Brotación", "Floración", "Cuajado", "Envero", "Maduración"],
            "Riesgo Principal": ["Oídio", "Black rot", "Oídio/Black rot", "Esca", "Botrytis"],
            "Acción Recomendada": [
                "Fungicida preventivo",
                "Fungicida sistémico",
                "Evaluación y aplicación según presión",
                "Monitoreo intensivo",
                "Aplicación pre-cosecha si es necesario"
            ]
        }

        calendar_df = pd.DataFrame(calendar_data)
        st.table(calendar_df)

# Ejecutar aplicación
if __name__ == "__main__":
    main()