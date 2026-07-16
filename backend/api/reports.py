import sys
import json
from pathlib import Path
from datetime import datetime
from typing import Any

from fastapi import APIRouter, HTTPException, Depends, status
from fastapi.responses import FileResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session

from backend.core.config import settings
from backend.core.security import get_current_user, TokenData
from backend.database.session import get_db
from backend.database.models import DiagnosticModel
from backend.repositories.diagnostic_repository import DiagnosticRepository
from backend.repositories.user_repository import UserRepository

BASE_DIR = Path(__file__).resolve().parent.parent.parent
if str(BASE_DIR) not in sys.path:
    sys.path.insert(0, str(BASE_DIR))

from src.services.prediction_service import DISEASE_CLASSES, DISEASE_INFO

router = APIRouter(prefix="/api/v1/reports", tags=["reports"])

REPORTS_OUTPUT_DIR = settings.STORAGE_DIR / "generated_reports"

# ---------------------------------------------------------------------------
# Paleta de marca VineGuard AI
# ---------------------------------------------------------------------------
COLOR_PRIMARY = "2E7D32"       # verde vid oscuro
COLOR_PRIMARY_LIGHT = "E8F5E9"  # verde muy claro (fondos)
COLOR_TEXT_MUTED = "6B7280"    # gris texto secundario
COLOR_RISK = {
    "high": "C62828",     # rojo
    "alto": "C62828",
    "medium": "F9A825",   # ambar
    "medio": "F9A825",
    "low": "2E7D32",      # verde
    "bajo": "2E7D32",
    "none": "2E7D32",
    "sano": "2E7D32",
}


def _risk_color(risk_level: str) -> str:
    return COLOR_RISK.get((risk_level or "").strip().lower(), COLOR_TEXT_MUTED)


def _safe_value(value: Any) -> str:
    if value is None:
        return "N/A"
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d %H:%M:%S")
    if isinstance(value, str):
        stripped = value.strip()
        return stripped or "N/A"
    return str(value)


def _coerce_float(value: Any) -> float | None:
    if isinstance(value, bool):
        return None
    if isinstance(value, (int, float)):
        return float(value)
    if isinstance(value, str):
        stripped = value.strip().replace("%", "")
        if not stripped:
            return None
        try:
            return float(stripped)
        except ValueError:
            return None
    return None


def _format_percentage(value: Any) -> str:
    number = _coerce_float(value)
    if number is None:
        return _safe_value(value)
    if 0 <= number <= 1:
        return f"{number:.2%}"
    if 1 < number <= 100:
        return f"{number:.2f}%"
    return f"{number:.4f}"


def _normalize_probabilities(probabilities: Any) -> tuple[list[tuple[str, Any]], str | None]:
    if probabilities is None:
        return [], None

    parsed = probabilities
    if isinstance(probabilities, str):
        stripped = probabilities.strip()
        if not stripped:
            return [], None
        try:
            parsed = json.loads(stripped)
        except json.JSONDecodeError:
            return [], _safe_value(probabilities)

    if isinstance(parsed, dict):
        items = []
        seen = set()
        for class_name in DISEASE_CLASSES:
            if class_name in parsed:
                items.append((class_name, parsed.get(class_name)))
                seen.add(class_name)
        for class_name, probability in parsed.items():
            if class_name not in seen:
                items.append((_safe_value(class_name), probability))
        return items, None

    if isinstance(parsed, list):
        items = []
        for index, probability in enumerate(parsed):
            class_name = DISEASE_CLASSES[index] if index < len(DISEASE_CLASSES) else f"Clase {index + 1}"
            items.append((class_name, probability))
        return items, None

    return [], _safe_value(parsed)


class ReportRequest(BaseModel):
    format: str = "docx"


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------
def _generate_docx_report(diagnosis: DiagnosticModel) -> Path:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx no instalado")

    REPORTS_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    def shade_cell(cell, hex_color):
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), hex_color)
        cell._tc.get_or_add_tcPr().append(shd)

    def set_cell_text(cell, text, *, bold=False, color=None, size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT):
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = alignment
        run = paragraph.add_run(_safe_value(text))
        run.font.bold = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor.from_string(color)

    def add_section_heading(doc, text):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(5)
        run = paragraph.add_run(text)
        run.font.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor.from_string(COLOR_PRIMARY)

    def add_info_table(doc, rows, *, highlighted_labels=None, risk_label=None, risk_fill_hex=None):
        highlighted_labels = highlighted_labels or set()
        table = doc.add_table(rows=0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = True
        for label, value in rows:
            row = table.add_row().cells
            set_cell_text(row[0], label, bold=True)
            set_cell_text(row[1], value)
            shade_cell(row[0], "F1F8F2")
            if label in highlighted_labels:
                shade_cell(row[0], COLOR_PRIMARY_LIGHT)
                shade_cell(row[1], COLOR_PRIMARY_LIGHT)
                set_cell_text(row[1], value, bold=True, color=COLOR_PRIMARY)
            if risk_label and label == risk_label and risk_fill_hex:
                shade_cell(row[1], risk_fill_hex)
                set_cell_text(row[1], value, bold=True, color="FFFFFF")
        return table

    result_key = diagnosis.result if isinstance(diagnosis.result, str) and diagnosis.result.strip() else None
    disease_info = DISEASE_INFO.get(result_key, {}) if result_key else {}
    predicted_class = _safe_value(diagnosis.result)
    risk_level = _safe_value(disease_info.get("risk_level"))
    risk_hex = _risk_color("" if risk_level == "N/A" else risk_level)
    probability_items, probability_fallback = _normalize_probabilities(diagnosis.probabilities)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("VineGuard AI")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(COLOR_PRIMARY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Reporte de Diagnóstico de Hoja de Vid")
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = RGBColor.from_string(COLOR_TEXT_MUTED)
    doc.add_paragraph()

    add_section_heading(doc, "Información General")
    add_info_table(doc, [
        ("ID del Diagnóstico", _safe_value(diagnosis.id)),
        ("Fecha", _safe_value(diagnosis.timestamp)),
        ("Archivo", _safe_value(diagnosis.filename)),
    ])
    doc.add_paragraph()

    add_section_heading(doc, "Resultado")
    add_info_table(doc, [
        ("Clase Predicha", predicted_class),
        ("Confianza", _format_percentage(diagnosis.confidence)),
        ("Modelo", _safe_value(diagnosis.model_used)),
    ], highlighted_labels={"Clase Predicha"})
    doc.add_paragraph()

    add_section_heading(doc, "Detalles de la Enfermedad")
    add_info_table(doc, [
        ("Nombre (ES)", _safe_value(disease_info.get("display_name_es"))),
        ("Nombre (EN)", _safe_value(disease_info.get("display_name_en"))),
        ("Nombre Científico", _safe_value(disease_info.get("scientific_name"))),
        ("Estado", _safe_value(disease_info.get("health_status"))),
        ("Nivel de Riesgo", risk_level.upper()),
    ], risk_label="Nivel de Riesgo", risk_fill_hex=risk_hex)
    doc.add_paragraph()

    add_section_heading(doc, "Distribución de Probabilidades")
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "Clase", bold=True, color="FFFFFF")
    set_cell_text(hdr[1], "Probabilidad", bold=True, color="FFFFFF")
    for cell in hdr:
        shade_cell(cell, COLOR_PRIMARY)

    if probability_items:
        for cls_name, probability in probability_items:
            row = table.add_row().cells
            set_cell_text(row[0], cls_name)
            set_cell_text(row[1], _format_percentage(probability))
            if cls_name == predicted_class:
                shade_cell(row[0], COLOR_PRIMARY_LIGHT)
                shade_cell(row[1], COLOR_PRIMARY_LIGHT)
                set_cell_text(row[0], cls_name, bold=True, color=COLOR_PRIMARY)
                set_cell_text(row[1], _format_percentage(probability), bold=True, color=COLOR_PRIMARY)
    else:
        row = table.add_row().cells
        if probability_fallback:
            set_cell_text(row[0], "Formato recibido")
            set_cell_text(row[1], probability_fallback)
        else:
            set_cell_text(row[0], "N/A")
            set_cell_text(row[1], "N/A")

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(18)
    footer_run = footer.add_run(
        "Este resultado es una estimación generada por IA y no reemplaza la evaluación "
        "de un ingeniero agrónomo o especialista fitosanitario."
    )
    footer_run.font.italic = True
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor.from_string(COLOR_TEXT_MUTED)

    filename = f"diagnostico_{diagnosis.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    output_path = REPORTS_OUTPUT_DIR / filename
    doc.save(str(output_path))
    return output_path


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------
def _generate_pdf_report(diagnosis: DiagnosticModel) -> Path:
    from xml.sax.saxutils import escape

    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
        HRFlowable,
        KeepTogether,
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER

    REPORTS_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    filename = f"diagnostico_{diagnosis.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    output_path = REPORTS_OUTPUT_DIR / filename

    result_key = diagnosis.result if isinstance(diagnosis.result, str) and diagnosis.result.strip() else None
    disease_info = DISEASE_INFO.get(result_key, {}) if result_key else {}
    predicted_class = _safe_value(diagnosis.result)
    risk_level = _safe_value(disease_info.get("risk_level"))
    risk_hex = _risk_color("" if risk_level == "N/A" else risk_level)
    probability_items, probability_fallback = _normalize_probabilities(diagnosis.probabilities)

    styles = getSampleStyleSheet()
    body_style = ParagraphStyle("VineBody", parent=styles["Normal"], fontSize=9, leading=12)
    title_style = ParagraphStyle(
        "VineTitle",
        parent=styles["Title"],
        textColor=colors.HexColor(f"#{COLOR_PRIMARY}"),
        fontSize=24,
        alignment=TA_CENTER,
        spaceAfter=6,
    )
    subtitle_style = ParagraphStyle(
        "VineSubtitle",
        parent=styles["Normal"],
        textColor=colors.HexColor(f"#{COLOR_TEXT_MUTED}"),
        fontSize=12,
        alignment=TA_CENTER,
        spaceAfter=10,
    )
    heading_style = ParagraphStyle(
        "VineHeading",
        parent=styles["Heading2"],
        textColor=colors.HexColor(f"#{COLOR_PRIMARY}"),
        fontSize=14,
        spaceBefore=10,
        spaceAfter=6,
    )
    footer_style = ParagraphStyle(
        "VineFooter",
        parent=styles["Normal"],
        textColor=colors.HexColor(f"#{COLOR_TEXT_MUTED}"),
        fontSize=8,
        alignment=TA_CENTER,
        leading=10,
    )

    def format_pdf_value(value, *, color_hex=None, bold=False):
        text = escape(_safe_value(value))
        if color_hex:
            text = f'<font color="#{color_hex}">{text}</font>'
        if bold:
            text = f"<b>{text}</b>"
        return Paragraph(text, body_style)

    def info_table(rows, *, highlighted_labels=None, risk_label=None, risk_fill_hex=None):
        highlighted_labels = highlighted_labels or set()
        data = [
            [
                Paragraph(f"<b>{escape(label)}</b>", body_style),
                format_pdf_value(value),
            ]
            for label, value in rows
        ]
        table = Table(data, colWidths=[5.1 * cm, 10.2 * cm], hAlign="LEFT")
        style_commands = [
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F1F8F2")),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#E0E0E0")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]
        for row_index, (label, value) in enumerate(rows):
            if label in highlighted_labels:
                style_commands.append(("BACKGROUND", (0, row_index), (-1, row_index), colors.HexColor(f"#{COLOR_PRIMARY_LIGHT}")))
                data[row_index][1] = format_pdf_value(value, color_hex=COLOR_PRIMARY, bold=True)
            if risk_label and label == risk_label and risk_fill_hex:
                style_commands.append(("BACKGROUND", (1, row_index), (1, row_index), colors.HexColor(f"#{risk_fill_hex}")))
                style_commands.append(("TEXTCOLOR", (1, row_index), (1, row_index), colors.white))
                data[row_index][1] = format_pdf_value(value, bold=True)
        table._cellvalues = data
        table.setStyle(TableStyle(style_commands))
        return table

    story = [
        Paragraph("VineGuard AI", title_style),
        Paragraph("Reporte de Diagnóstico de Hoja de Vid", subtitle_style),
        HRFlowable(width="100%", color=colors.HexColor(f"#{COLOR_PRIMARY}"), thickness=1.5),
        Spacer(1, 12),
        Paragraph("Información General", heading_style),
        info_table([
            ("ID del Diagnóstico", _safe_value(diagnosis.id)),
            ("Fecha", _safe_value(diagnosis.timestamp)),
            ("Archivo", _safe_value(diagnosis.filename)),
        ]),
        Spacer(1, 10),
        Paragraph("Resultado", heading_style),
        info_table([
            ("Clase Predicha", predicted_class),
            ("Confianza", _format_percentage(diagnosis.confidence)),
            ("Modelo", _safe_value(diagnosis.model_used)),
        ], highlighted_labels={"Clase Predicha"}),
        Spacer(1, 10),
        Paragraph("Detalles de la Enfermedad", heading_style),
        info_table([
            ("Nombre (ES)", _safe_value(disease_info.get("display_name_es"))),
            ("Nombre (EN)", _safe_value(disease_info.get("display_name_en"))),
            ("Nombre Científico", _safe_value(disease_info.get("scientific_name"))),
            ("Estado", _safe_value(disease_info.get("health_status"))),
            ("Nivel de Riesgo", risk_level.upper()),
        ], risk_label="Nivel de Riesgo", risk_fill_hex=risk_hex),
        Spacer(1, 10),
        Paragraph("Distribución de Probabilidades", heading_style),
    ]

    probability_data = [["Clase", "Probabilidad"]]
    if probability_items:
        probability_data.extend([[cls_name, _format_percentage(probability)] for cls_name, probability in probability_items])
    elif probability_fallback:
        probability_data.append(["Formato recibido", probability_fallback])
    else:
        probability_data.append(["N/A", "N/A"])

    probabilities_table = Table(probability_data, colWidths=[10.2 * cm, 5.1 * cm], repeatRows=1, hAlign="LEFT")
    probability_style = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{COLOR_PRIMARY}")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#E0E0E0")),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]
    if probability_items:
        for row_index, (class_name, _) in enumerate(probability_items, start=1):
            if class_name == predicted_class:
                probability_style.append(("BACKGROUND", (0, row_index), (-1, row_index), colors.HexColor(f"#{COLOR_PRIMARY_LIGHT}")))
                probability_style.append(("TEXTCOLOR", (0, row_index), (-1, row_index), colors.HexColor(f"#{COLOR_PRIMARY}")))
                probability_style.append(("FONTNAME", (0, row_index), (-1, row_index), "Helvetica-Bold"))
    probabilities_table.setStyle(TableStyle(probability_style))
    story.append(probabilities_table)

    story.append(KeepTogether([
        Spacer(1, 22),
        HRFlowable(width="100%", color=colors.HexColor("#D6D6D6"), thickness=0.5),
        Spacer(1, 8),
        Paragraph(
            "Este resultado es una estimación generada por IA y no reemplaza la evaluación "
            "de un ingeniero agrónomo o especialista fitosanitario.",
            footer_style,
        ),
    ]))

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        topMargin=2 * cm,
        bottomMargin=2.8 * cm,
        leftMargin=1.8 * cm,
        rightMargin=1.8 * cm,
    )
    doc.build(story)
    return output_path


# ---------------------------------------------------------------------------
# XLSX
# ---------------------------------------------------------------------------
def _generate_excel_report(diagnosis: DiagnosticModel) -> Path:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    REPORTS_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    filename = f"diagnostico_{diagnosis.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    output_path = REPORTS_OUTPUT_DIR / filename

    result_key = diagnosis.result if isinstance(diagnosis.result, str) and diagnosis.result.strip() else None
    disease_info = DISEASE_INFO.get(result_key, {}) if result_key else {}
    predicted_class = _safe_value(diagnosis.result)
    risk_level = _safe_value(disease_info.get("risk_level"))
    risk_hex = _risk_color("" if risk_level == "N/A" else risk_level)
    probability_items, probability_fallback = _normalize_probabilities(diagnosis.probabilities)

    primary_fill = PatternFill(start_color=COLOR_PRIMARY, end_color=COLOR_PRIMARY, fill_type="solid")
    light_fill = PatternFill(start_color=COLOR_PRIMARY_LIGHT, end_color=COLOR_PRIMARY_LIGHT, fill_type="solid")
    label_fill = PatternFill(start_color="F1F8F2", end_color="F1F8F2", fill_type="solid")
    risk_fill = PatternFill(start_color=risk_hex, end_color=risk_hex, fill_type="solid")
    thin_side = Side(style="thin", color="E0E0E0")
    thin_border = Border(left=thin_side, right=thin_side, top=thin_side, bottom=thin_side)
    white_bold = Font(bold=True, color="FFFFFF", size=11)
    section_font = Font(bold=True, size=12, color=COLOR_PRIMARY)
    label_font = Font(bold=True, size=10)
    highlight_font = Font(bold=True, size=11, color=COLOR_PRIMARY)
    risk_font = Font(bold=True, size=11, color="FFFFFF")
    muted_font = Font(italic=True, size=9, color=COLOR_TEXT_MUTED)

    wb = Workbook()
    ws = wb.active
    ws.title = "Diagnóstico"
    ws.sheet_view.showGridLines = False
    ws.column_dimensions["A"].width = 28
    ws.column_dimensions["B"].width = 42

    def style_value_cell(cell, *, font=None, fill=None):
        cell.border = thin_border
        cell.alignment = Alignment(vertical="top", wrap_text=True)
        if font:
            cell.font = font
        if fill:
            cell.fill = fill

    def write_title():
        ws.merge_cells("A1:B1")
        c = ws["A1"]
        c.value = "VineGuard AI"
        c.font = Font(bold=True, size=18, color="FFFFFF")
        c.fill = primary_fill
        c.alignment = Alignment(horizontal="center", vertical="center")

        ws.merge_cells("A2:B2")
        c2 = ws["A2"]
        c2.value = "Reporte de Diagnóstico de Hoja de Vid"
        c2.font = Font(italic=True, size=11, color=COLOR_TEXT_MUTED)
        c2.fill = light_fill
        c2.alignment = Alignment(horizontal="center", vertical="center")
        ws.append([])

    def write_section(title):
        ws.append([title])
        cell = ws.cell(row=ws.max_row, column=1)
        cell.font = section_font
        cell.alignment = Alignment(horizontal="left")
        ws.merge_cells(start_row=ws.max_row, start_column=1, end_row=ws.max_row, end_column=2)

    def write_row(label, value, *, value_font=None, value_fill=None, highlight=False):
        ws.append([label, value])
        row_index = ws.max_row
        label_cell = ws.cell(row=row_index, column=1)
        value_cell = ws.cell(row=row_index, column=2)

        label_cell.font = label_font
        label_cell.fill = light_fill if highlight else label_fill
        label_cell.border = thin_border
        label_cell.alignment = Alignment(vertical="top", wrap_text=True)

        style_value_cell(value_cell, font=value_font, fill=value_fill or (light_fill if highlight else None))
        return row_index

    write_title()
    write_section("Información General")
    write_row("ID del Diagnóstico", _safe_value(diagnosis.id))
    write_row("Fecha", _safe_value(diagnosis.timestamp))
    write_row("Archivo", _safe_value(diagnosis.filename))
    ws.append([])

    write_section("Resultado")
    write_row("Clase Predicha", predicted_class, value_font=highlight_font, highlight=True)
    write_row("Confianza", _format_percentage(diagnosis.confidence))
    write_row("Modelo", _safe_value(diagnosis.model_used))
    ws.append([])

    write_section("Detalles de la Enfermedad")
    write_row("Nombre (ES)", _safe_value(disease_info.get("display_name_es")))
    write_row("Nombre (EN)", _safe_value(disease_info.get("display_name_en")))
    write_row("Nombre Científico", _safe_value(disease_info.get("scientific_name")))
    write_row("Estado", _safe_value(disease_info.get("health_status")))
    write_row("Nivel de Riesgo", risk_level.upper(), value_font=risk_font, value_fill=risk_fill)
    ws.append([])

    write_section("Distribución de Probabilidades")
    header_row = ws.max_row + 1
    ws.append(["Clase", "Probabilidad"])
    for col in (1, 2):
        cell = ws.cell(row=header_row, column=col)
        cell.fill = primary_fill
        cell.font = white_bold
        cell.border = thin_border
        cell.alignment = Alignment(horizontal="left", vertical="center")

    if probability_items:
        for cls_name, probability in probability_items:
            ws.append([cls_name, _format_percentage(probability)])
            row_index = ws.max_row
            for col in (1, 2):
                cell = ws.cell(row=row_index, column=col)
                cell.border = thin_border
                cell.alignment = Alignment(vertical="top", wrap_text=True)
                if cls_name == predicted_class:
                    cell.fill = light_fill
                    cell.font = highlight_font
    elif probability_fallback:
        write_row("Formato recibido", probability_fallback)
    else:
        write_row("N/A", "N/A")

    ws.append([])
    footer_row = ws.max_row + 1
    ws.merge_cells(start_row=footer_row, start_column=1, end_row=footer_row, end_column=2)
    footer_cell = ws.cell(row=footer_row, column=1)
    footer_cell.value = (
        "Este resultado es una estimación generada por IA y no reemplaza la evaluación "
        "de un ingeniero agrónomo o especialista fitosanitario."
    )
    footer_cell.font = muted_font
    footer_cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    wb.save(str(output_path))
    return output_path


# ---------------------------------------------------------------------------
# Rutas
# ---------------------------------------------------------------------------
@router.get("")
def list_reports(current_user: TokenData = Depends(get_current_user)):
    REPORTS_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    files = []
    for f in REPORTS_OUTPUT_DIR.iterdir():
        if f.is_file() and f.suffix in {".docx", ".pdf", ".xlsx"}:
            files.append({
                "id": f.stem,
                "filename": f.name,
                "size_bytes": f.stat().st_size,
                "created_at": datetime.fromtimestamp(f.stat().st_ctime).isoformat(),
            })
    return {"reports": sorted(files, key=lambda x: x["created_at"], reverse=True)}


@router.post("/diagnosis/{diagnosis_id}")
def generate_report(
    diagnosis_id: int,
    payload: ReportRequest = ReportRequest(),
    db: Session = Depends(get_db),
    current_user: TokenData = Depends(get_current_user),
):
    diag_repo = DiagnosticRepository(db)
    diag = diag_repo.get_by_id(diagnosis_id)
    if not diag:
        raise HTTPException(status_code=404, detail="Diagnóstico no encontrado")

    user_repo = UserRepository(db)
    user = user_repo.get_by_username(current_user.sub)
    if not user:
        raise HTTPException(status_code=404, detail="Usuario no encontrado")

    if current_user.role != "admin" and diag.user_id != user.id:
        raise HTTPException(status_code=403, detail="No autorizado")

    fmt = payload.format.lower()
    if fmt == "pdf":
        output_path = _generate_pdf_report(diag)
    elif fmt == "xlsx":
        output_path = _generate_excel_report(diag)
    else:
        output_path = _generate_docx_report(diag)

    return {
        "message": "Reporte generado exitosamente",
        "filename": output_path.name,
        "path": str(output_path),
        "download_url": f"/api/v1/reports/{output_path.stem}/download",
    }


@router.get("/{report_id}/download")
def download_report(report_id: str, current_user: TokenData = Depends(get_current_user)):
    REPORTS_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    for f in REPORTS_OUTPUT_DIR.iterdir():
        if f.stem == report_id and f.is_file():
            media_types = {
                ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ".pdf": "application/pdf",
                ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            }
            return FileResponse(
                path=str(f),
                filename=f.name,
                media_type=media_types.get(f.suffix, "application/octet-stream"),
            )
    raise HTTPException(status_code=404, detail="Reporte no encontrado")
